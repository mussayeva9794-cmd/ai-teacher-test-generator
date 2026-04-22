"""Streamlit web app for generating, editing, practicing, and analyzing teacher tests."""

from __future__ import annotations

from copy import deepcopy
from collections import defaultdict
from datetime import datetime, timedelta
from functools import lru_cache
from io import BytesIO
import hashlib
import json
import os
import random
from threading import Thread
from typing import Any
from uuid import uuid4
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Pt
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
import streamlit as st
import streamlit.components.v1 as components

from ai_generator import generate_test
from analytics import (
    aggregate_attempt_history,
    build_gradebook_rows,
    build_topic_progress_rows,
    classify_risk,
    detect_suspicious_attempts,
    grade_attempt,
)
from cloud_sync import get_cloud_status, is_cloud_enabled, sync_attempt_result, sync_history_record, sync_question_bank_item
from document_loader import SUPPORTED_EXTENSIONS, extract_text_from_uploaded_file
from quality import analyze_test_quality
from storage import (
    attempt_submission_exists,
    authenticate_local_user,
    create_share_link,
    create_student_group,
    create_local_user,
    delete_attempt_result,
    count_share_attempts,
    count_share_attempts_for_student_key,
    delete_student_draft,
    get_plan_status,
    import_group_students,
    initialize_database,
    is_valid_email,
    list_audit_logs,
    list_api_error_logs,
    list_attempt_results,
    list_group_students,
    list_student_groups,
    list_question_bank,
    list_share_links,
    list_test_library,
    list_test_history,
    list_usage_events,
    log_audit_event,
    migrate_local_data_to_cloud,
    load_attempt_result,
    load_student_draft,
    load_share_link,
    load_question_bank_item,
    load_latest_test_record,
    load_test_record,
    log_api_error,
    record_usage_event,
    save_attempt_result,
    save_group_student,
    save_student_draft,
    save_question_bank_item,
    save_test_record,
    set_share_link_status,
    set_test_archived,
    set_test_favorite,
    upsert_autosave_record,
    update_attempt_result,
)
from variants import build_all_variants


st.set_page_config(
    page_title="AI Teacher Test Generator",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="collapsed",
)


DIFFICULTY_OPTIONS = {
    "easy": "Easy",
    "medium": "Medium",
    "hard": "Hard",
}

TEST_TYPE_OPTIONS = {
    "multiple_choice": "Multiple Choice",
    "true_false": "True / False",
    "short_answer": "Short Answer",
    "matching": "Matching",
}

LANGUAGE_OPTIONS = {
    "english": "English",
    "russian": "Russian",
    "kazakh": "Kazakh",
}

GRADE_LEVEL_OPTIONS = [
    "Primary school",
    "5-6 grade",
    "7-9 grade",
    "10-11 grade",
    "College",
    "University",
]

LESSON_STAGE_OPTIONS = [
    "Introduction",
    "Practice",
    "Revision",
    "Assessment",
    "Homework",
]

ASSESSMENT_PURPOSE_OPTIONS = [
    "Formative assessment",
    "Quick quiz",
    "Homework check",
    "Exam preparation",
    "Independent practice",
]


def default_guest_user() -> dict[str, Any]:
    """Return the default guest profile."""
    return {
        "email": "guest@local",
        "display_name": "Guest Teacher",
        "role": "teacher",
        "is_guest": True,
    }


def initialize_state() -> None:
    """Create all session state keys used by the app."""
    initialize_database()
    defaults = {
        "generated_test": None,
        "generated_variants": {},
        "generated_topic": "",
        "editor_version": 0,
        "test_metadata": {},
        "history_notice": "",
        "source_preview": "",
        "source_stats": {},
        "quality_report": None,
        "current_user": default_guest_user(),
        "last_attempt": None,
        "question_bank_notice": "",
        "share_notice": "",
        "public_app_url": os.getenv("PUBLIC_APP_URL", "http://localhost:8501"),
        "last_autosave_signature": "",
        "last_created_share_url": "",
        "active_flow_step": "Review",
        "student_draft_loaded_for": "",
        "student_submission_confirmed": False,
        "onboarding_dismissed": False,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def get_current_user() -> dict[str, Any]:
    """Return the active user profile."""
    return st.session_state.current_user


def get_owner_email() -> str:
    """Return the owner email used for local records."""
    return get_current_user()["email"]


def get_current_test_uid() -> str:
    """Return the stable identifier of the current test pack."""
    return str(st.session_state.test_metadata.get("test_uid", "")).strip()


def normalize_tag_text(raw_value: str) -> str:
    """Normalize one subject tag."""
    return " ".join(raw_value.strip().split())


def parse_subject_tags(raw_value: str) -> list[str]:
    """Split comma-separated subject tags into a stable list."""
    tags: list[str] = []
    for part in str(raw_value).split(","):
        cleaned = normalize_tag_text(part)
        if cleaned and cleaned.lower() not in {item.lower() for item in tags}:
            tags.append(cleaned)
    return tags


def format_subject_tags(raw_value: str) -> str:
    """Return normalized comma-separated subject tags."""
    return ", ".join(parse_subject_tags(raw_value))


def build_payload_signature(test_data: dict[str, Any], metadata: dict[str, Any]) -> str:
    """Build a stable signature for auto-save detection."""
    raw = json.dumps({"test_data": test_data, "metadata": metadata}, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def autosave_current_test(test_data: dict[str, Any], metadata: dict[str, Any]) -> None:
    """Autosave the current test only when its content changes."""
    if not metadata.get("test_uid"):
        return
    signature = build_payload_signature(test_data, metadata)
    if st.session_state.get("last_autosave_signature") == signature:
        return

    upsert_autosave_record(
        test_uid=metadata.get("test_uid", ""),
        title=test_data.get("title", "Generated Test"),
        topic=metadata.get("topic", ""),
        language=metadata.get("language", "english"),
        difficulty=metadata.get("difficulty", "medium"),
        test_type=metadata.get("test_type", "multiple_choice"),
        grade_level=metadata.get("grade_level", ""),
        assessment_purpose=metadata.get("assessment_purpose", ""),
        owner_email=get_owner_email(),
        source_kind=metadata.get("source_kind", "topic"),
        source_name=metadata.get("source_name", ""),
        subject_tags=format_subject_tags(metadata.get("subject_tags", "")),
        is_favorite=bool(metadata.get("is_favorite", False)),
        payload={"test_data": test_data, "variants": st.session_state.get("generated_variants", {}), "metadata": metadata},
    )
    st.session_state.last_autosave_signature = signature


def load_test_payload(payload: dict[str, Any]) -> None:
    """Load a test payload into the active workspace."""
    st.session_state.generated_test = payload["test_data"]
    st.session_state.generated_variants = payload.get("variants", {})
    st.session_state.test_metadata = payload["metadata"]
    if not st.session_state.test_metadata.get("test_uid"):
        st.session_state.test_metadata["test_uid"] = uuid4().hex
    st.session_state.test_metadata.setdefault("subject_tags", "")
    st.session_state.test_metadata.setdefault("is_favorite", False)
    st.session_state.generated_topic = payload["metadata"].get("topic", "")
    st.session_state.editor_version += 1
    st.session_state.source_preview = ""
    st.session_state.source_stats = {}
    st.session_state.quality_report = analyze_test_quality(
        payload["test_data"],
        expected_count=payload["metadata"].get("question_count"),
    )
    st.session_state.active_flow_step = "Review"
    st.session_state.last_autosave_signature = build_payload_signature(
        st.session_state.generated_test,
        st.session_state.test_metadata,
    )


def clear_workspace() -> None:
    """Clear the active workspace."""
    st.session_state.generated_test = None
    st.session_state.generated_variants = {}
    st.session_state.generated_topic = ""
    st.session_state.test_metadata = {}
    st.session_state.quality_report = None
    st.session_state.last_attempt = None
    st.session_state.active_flow_step = "Review"
    st.session_state.last_autosave_signature = ""


def duplicate_current_test() -> None:
    """Duplicate the current test into a fresh working copy."""
    if not st.session_state.generated_test or not st.session_state.test_metadata:
        return
    new_uid = uuid4().hex
    duplicated_test = deepcopy(st.session_state.generated_test)
    duplicated_variants = deepcopy(st.session_state.get("generated_variants", {}))
    for variant_name, variant_data in duplicated_variants.items():
        variant_data["test_uid"] = new_uid
        variant_data["variant_name"] = variant_name
    duplicated_test["test_uid"] = new_uid
    new_metadata = deepcopy(st.session_state.test_metadata)
    new_metadata["test_uid"] = new_uid
    new_metadata["topic"] = f"{new_metadata.get('topic', '')} copy".strip()
    new_metadata["is_favorite"] = False
    load_test_payload(
        {
            "test_data": duplicated_test,
            "variants": duplicated_variants,
            "metadata": new_metadata,
        }
    )


def archive_current_test() -> None:
    """Archive the current active test."""
    current_uid = get_current_test_uid()
    if not current_uid:
        return
    set_test_archived(current_uid, get_owner_email(), True)
    st.session_state.history_notice = "Current test was archived."
    clear_workspace()


def copy_share_link_value(share_url: str) -> None:
    """Store a share URL so it can be shown in the UI."""
    st.session_state.last_created_share_url = share_url
    st.session_state.share_notice = "Share link is ready to copy."


def get_status_label() -> str:
    """Return a user-friendly workspace status label."""
    if st.session_state.generated_test is None:
        return "No test"
    attempts = list_attempt_results(limit=1, owner_email=get_owner_email(), test_uid=get_current_test_uid())
    if attempts:
        return "Has responses"
    return "Test ready"


def get_owner_roster() -> list[dict[str, Any]]:
    """Return imported roster rows for the active teacher."""
    return list_group_students(get_owner_email())


def get_current_plan_status() -> dict[str, Any]:
    """Return current teacher plan and usage state."""
    return get_plan_status(get_owner_email())


def can_use_generation() -> tuple[bool, str]:
    """Check whether the current teacher can generate another test."""
    current_user = get_current_user()
    if current_user.get("role") != "teacher":
        return False, "Only teacher accounts can generate new tests."
    plan = get_current_plan_status()
    limit = int(plan["limits"].get("monthly_generations", 0))
    used = int(plan["usage"].get("monthly_generations", 0))
    if limit and used >= limit:
        return False, f"Your {plan['plan_name']} plan reached the monthly generation limit ({used}/{limit})."
    active_tests = len(list_test_library(owner_email=get_owner_email(), include_archived=False))
    active_limit = int(plan["limits"].get("active_tests", 0))
    if active_limit and active_tests >= active_limit:
        return False, f"Your {plan['plan_name']} plan already has the maximum number of active tests ({active_tests}/{active_limit}). Archive old tests before creating a new one."
    return True, ""


def can_add_students(extra_count: int = 1) -> tuple[bool, str]:
    """Check whether the current roster can grow by the requested amount."""
    plan = get_current_plan_status()
    current_students = len(get_owner_roster())
    student_limit = int(plan["limits"].get("students", 0))
    if student_limit and current_students + extra_count > student_limit:
        return False, f"Adding {extra_count} students would exceed the {plan['plan_name']} plan limit ({current_students}/{student_limit})."
    return True, ""


def log_event(event_type: str, target_type: str = "", target_id: str = "", details: dict[str, Any] | None = None) -> None:
    """Write one audit trail event for the current user."""
    current_user = get_current_user()
    log_audit_event(
        current_user.get("email", "guest@local"),
        current_user.get("role", ""),
        event_type,
        target_type,
        target_id,
        details or {},
    )


def friendly_generation_error_message(error: Exception) -> str:
    """Translate raw generation failures into actionable teacher-facing text."""
    message = str(error).strip()
    lower = message.lower()
    if "api key" in lower or "authentication" in lower:
        return "The AI provider rejected the request. Check the configured API key in local .env or Streamlit Secrets."
    if "rate limit" in lower or "quota" in lower:
        return "The AI provider is temporarily overloaded or out of quota. Wait a moment and try again."
    if "json" in lower or "schema" in lower:
        return "The AI returned an invalid structured response. Try again or slightly simplify the topic and settings."
    if "network" in lower or "connection" in lower or "timed out" in lower:
        return "The app could not reach the AI provider. Check internet access and cloud availability."
    return message or "The test could not be generated. Please review the settings and try again."


def build_answer_signature(responses: dict[str, Any]) -> str:
    """Build a stable response fingerprint for anti-cheat analytics."""
    normalized = json.dumps(responses, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def render_onboarding_panel() -> None:
    """Show a compact onboarding checklist for new teachers."""
    if st.session_state.get("onboarding_dismissed"):
        return
    with st.expander("Quick start for teachers", expanded=st.session_state.generated_test is None):
        st.markdown(
            """
1. Create a test from a topic or source file.
2. Review and adjust the questions in **Review**.
3. Create a protected student link in **Share**.
4. Track outcomes in **Analyze** with gradebook, topic progress, and suspicious-attempt flags.
            """
        )
        info_cols = st.columns(3)
        info_cols[0].info("Tip: import your student roster first to make the gradebook more useful.")
        info_cols[1].info("Tip: use one-question mode and a timer for higher-stakes work.")
        info_cols[2].info("Tip: export a backup before major edits or live demos.")
        if st.button("Hide onboarding", key="hide_onboarding", use_container_width=True):
            st.session_state.onboarding_dismissed = True
            st.rerun()

def render_theme() -> None:
    """Inject the custom visual theme for the app."""
    st.markdown(
        """
        <style>
        :root {
            --wine: #951122;
            --ink: #100C08;
            --mist: #f7f1ec;
            --mist-2: #dccfc5;
            --line: rgba(149, 17, 34, 0.16);
            --line-strong: rgba(149, 17, 34, 0.34);
            --text-soft: #c8b8ab;
            --shadow: 0 24px 60px rgba(0, 0, 0, 0.22);
        }

        .stApp {
            background:
                radial-gradient(circle at 12% 10%, rgba(149, 17, 34, 0.12), transparent 28%),
                radial-gradient(circle at 88% 16%, rgba(149, 17, 34, 0.10), transparent 24%),
                radial-gradient(circle at 50% 100%, rgba(149, 17, 34, 0.07), transparent 28%),
                linear-gradient(180deg, #0d0907 0%, #120c09 54%, #0f0b08 100%);
            color: var(--mist);
            font-family: "Avenir Next", "SF Pro Display", "Segoe UI", sans-serif;
        }

        .block-container {
            padding-top: 1rem;
            padding-bottom: 3rem;
            max-width: 1120px;
        }

        [data-testid="stSidebar"] {
            background: linear-gradient(180deg, rgba(14, 10, 8, 0.94) 0%, rgba(19, 13, 10, 0.94) 100%);
            border-right: 1px solid rgba(255, 255, 255, 0.04);
            backdrop-filter: blur(22px);
        }

        h1, h2, h3, h4,
        [data-testid="stMarkdownContainer"] p,
        [data-testid="stMarkdownContainer"] li,
        [data-testid="stMetricLabel"],
        [data-testid="stMetricValue"],
        label {
            color: var(--mist);
        }

        .hero-shell {
            position: relative;
            overflow: hidden;
            border: 1px solid rgba(255, 255, 255, 0.06);
            border-radius: 26px;
            padding: 1.5rem 1.55rem 1.45rem;
            margin-bottom: 1rem;
            background:
                linear-gradient(140deg, rgba(149, 17, 34, 0.22), rgba(255,255,255,0.015) 46%, rgba(255,255,255,0.03) 100%),
                rgba(17, 12, 9, 0.78);
            box-shadow: var(--shadow);
            backdrop-filter: blur(20px);
        }

        .hero-shell::before {
            content: "";
            position: absolute;
            inset: 0;
            pointer-events: none;
            background-image:
                url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='920' height='720' viewBox='0 0 920 720' fill='none'%3E%3Cpath d='M124 156C238 86 336 88 420 154V530C336 466 238 466 124 530V156Z' stroke='%23ffffff' stroke-opacity='0.07' stroke-width='2.1'/%3E%3Cpath d='M420 154C502 88 600 86 716 156V530C600 466 502 466 420 530V154Z' stroke='%23ffffff' stroke-opacity='0.07' stroke-width='2.1'/%3E%3Ccircle cx='700' cy='170' r='70' stroke='%23ffffff' stroke-opacity='0.06' stroke-width='2'/%3E%3Ccircle cx='700' cy='170' r='18' fill='%23951122' fill-opacity='0.14'/%3E%3C/svg%3E");
            background-repeat: no-repeat;
            background-position: right -90px top 28px;
            background-size: min(42vw, 460px);
            opacity: 0.9;
        }

        .hero-shell::after {
            content: "";
            position: absolute;
            top: -64px;
            right: -28px;
            width: 220px;
            height: 220px;
            border-radius: 999px;
            background: radial-gradient(circle, rgba(149, 17, 34, 0.24), transparent 68%);
            filter: blur(6px);
        }

        .hero-kicker {
            position: relative;
            z-index: 1;
            font-size: 0.66rem;
            letter-spacing: 0.18em;
            text-transform: uppercase;
            color: #e7c9c1;
            margin-bottom: 0.55rem;
        }

        .hero-title {
            position: relative;
            z-index: 1;
            font-size: 2.35rem;
            line-height: 1.02;
            font-weight: 780;
            letter-spacing: -0.04em;
            margin: 0 0 0.6rem 0;
            color: #fffaf7;
        }

        .hero-copy {
            position: relative;
            z-index: 1;
            max-width: 620px;
            margin: 0;
            font-size: 0.98rem;
            line-height: 1.65;
            color: #d8c8bc;
        }

        .student-shell {
            margin-top: 0.3rem;
            background:
                linear-gradient(140deg, rgba(149, 17, 34, 0.20), rgba(255,255,255,0.015) 50%, rgba(255,255,255,0.02) 100%),
                rgba(17, 12, 9, 0.8);
        }

        .section-card {
            border: 1px solid rgba(255, 255, 255, 0.06);
            border-radius: 22px;
            padding: 0.95rem 1rem 1.02rem;
            background: linear-gradient(180deg, rgba(255,255,255,0.028), rgba(255,255,255,0.018));
            box-shadow: 0 12px 34px rgba(0, 0, 0, 0.16);
            backdrop-filter: blur(18px);
            margin-bottom: 0.95rem;
        }

        .section-label {
            font-size: 0.64rem;
            text-transform: uppercase;
            letter-spacing: 0.18em;
            color: #cfb0a6;
            margin-bottom: 0.7rem;
        }

        .workspace-hint {
            border: 1px dashed rgba(232, 221, 212, 0.14);
            border-radius: 24px;
            padding: 1.35rem 1.3rem;
            background: linear-gradient(180deg, rgba(255,255,255,0.03), rgba(255,255,255,0.015));
            margin-top: 0.85rem;
            backdrop-filter: blur(14px);
        }

        .workspace-hint strong {
            color: #fff6f1;
        }

        [data-testid="stTabs"] [role="tablist"] {
            gap: 0.35rem;
            padding: 0.14rem;
            background: rgba(255,255,255,0.015);
            border: 1px solid rgba(255,255,255,0.06);
            border-radius: 16px;
            backdrop-filter: blur(12px);
        }

        [data-testid="stTabs"] [role="tab"] {
            border-radius: 12px;
            padding: 0.48rem 0.9rem;
            color: #d9ccc2;
            background: transparent;
            transition: all 180ms ease;
            border: 1px solid transparent;
        }

        [data-testid="stTabs"] [aria-selected="true"] {
            background: linear-gradient(135deg, rgba(149, 17, 34, 0.9), rgba(121, 16, 30, 0.94));
            color: white;
            box-shadow: 0 8px 24px rgba(149, 17, 34, 0.18);
        }

        div[data-testid="stMetric"] {
            background: rgba(255,255,255,0.022);
            border: 1px solid rgba(255,255,255,0.06);
            border-radius: 18px;
            padding: 0.78rem 0.88rem;
            box-shadow: inset 0 1px 0 rgba(255,255,255,0.03);
        }

        div[data-baseweb="input"] > div,
        div[data-baseweb="select"] > div,
        .stTextArea textarea {
            background: rgba(255,255,255,0.035) !important;
            border-color: rgba(255,255,255,0.08) !important;
            border-radius: 16px !important;
            color: #fff8f5 !important;
            box-shadow: none !important;
            transition: border-color 180ms ease, background 180ms ease, transform 180ms ease !important;
        }

        div[data-baseweb="input"] > div:focus-within,
        div[data-baseweb="select"] > div:focus-within,
        .stTextArea textarea:focus {
            border-color: var(--line-strong) !important;
            background: rgba(255,255,255,0.05) !important;
        }

        .stSlider [data-baseweb="slider"] [role="slider"] {
            background-color: var(--wine);
            box-shadow: 0 0 0 6px rgba(149, 17, 34, 0.12);
        }

        .stButton > button,
        .stDownloadButton > button {
            border-radius: 16px;
            border: 1px solid rgba(255,255,255,0.08);
            background: linear-gradient(135deg, var(--wine), #6f0e1d);
            color: #fff7f4;
            min-height: 2.8rem;
            box-shadow: 0 14px 26px rgba(149, 17, 34, 0.16);
            transition: transform 180ms ease, box-shadow 180ms ease, border-color 180ms ease, filter 180ms ease;
        }

        .stButton > button:hover,
        .stDownloadButton > button:hover {
            border-color: rgba(255,255,255,0.18);
            background: linear-gradient(135deg, #a91529, #7e1421);
            transform: translateY(-2px) scale(1.01);
            box-shadow: 0 18px 32px rgba(149, 17, 34, 0.24);
            filter: saturate(1.06);
        }

        .stButton > button:active,
        .stDownloadButton > button:active {
            transform: translateY(0) scale(0.988);
        }

        .stExpander {
            border: 1px solid rgba(255,255,255,0.06) !important;
            border-radius: 18px !important;
            background: rgba(255,255,255,0.018);
            backdrop-filter: blur(12px);
        }

        [data-testid="stDataFrame"], [data-testid="stTable"] {
            border-radius: 18px;
            overflow: hidden;
            border: 1px solid rgba(255,255,255,0.06);
            background: rgba(255,255,255,0.02);
        }

        [data-testid="stDataFrame"] * {
            border-color: rgba(255,255,255,0.05) !important;
        }

        [data-testid="stPopover"] {
            backdrop-filter: blur(18px);
        }

        .stAlert {
            border-radius: 18px;
            border: 1px solid rgba(255,255,255,0.06);
        }

        #MainMenu,
        footer,
        .stAppDeployButton {
            visibility: hidden;
        }

        @media (max-width: 900px) {
            .block-container {
                padding-top: 0.8rem;
                padding-left: 0.8rem;
                padding-right: 0.8rem;
            }

            .hero-title {
                font-size: 1.7rem;
            }

            .hero-shell,
            .section-card {
                border-radius: 16px;
                padding: 0.95rem;
            }

            .hero-shell::before {
                background-position: right -100px top 36px;
                background-size: 90vw;
                opacity: 0.42;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_header() -> None:
    """Render the app title and intro text."""
    st.markdown(
        """
        <div class="hero-shell">
            <div class="hero-kicker">Teacher Studio</div>
            <div class="hero-title">AI Teacher Test Generator</div>
            <p class="hero-copy">
                A quieter workspace for building classroom tests, guiding students, and reading the signal behind their results.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_project_explainers() -> None:
    """Render short sections that strengthen the pedagogical and defense narrative."""
    with st.sidebar.expander("Project Notes", expanded=False):
        st.markdown("**Pedagogical value**")
        st.markdown(
            """
This system is designed for teachers and future informatics teachers.

- It supports formative assessment, revision, homework checking, and exam preparation.
- It adapts questions to grade level, lesson stage, and learning objective.
- It helps a teacher save time while still keeping full control through manual editing.
- It supports multilingual classroom practice in English, Russian, and Kazakh.
"""
        )
        st.markdown("**Cloud and platform**")
        st.markdown(
            """
- The project uses a real cloud AI service through Groq API.
- It supports optional Supabase cloud backup for history, attempts, and question bank data.
- Local profiles and history work even without cloud credentials.
- Student mode and analytics turn the application into a practical EdTech workflow, not just a text generator.
"""
        )


def open_section(title: str) -> None:
    """Open a styled visual section wrapper."""
    st.markdown(
        f"""
        <div class="section-card">
            <div class="section-label">{title}</div>
        """,
        unsafe_allow_html=True,
    )


def close_section() -> None:
    """Close a styled visual section wrapper."""
    st.markdown("</div>", unsafe_allow_html=True)


def render_profile_sidebar() -> None:
    """Render authentication and profile controls."""
    current_user = get_current_user()
    with st.sidebar.expander("Account", expanded=False):
        st.caption(f"Signed in as: {current_user['display_name']} ({current_user['role']})")
        if current_user.get("is_guest"):
            st.info("You are using guest mode. Create a local profile to keep personal history and question bank data.")

        sign_in_tab, sign_up_tab, guest_tab = st.tabs(["Sign In", "Sign Up", "Guest"])

        with sign_in_tab:
            with st.form("sign_in_form", clear_on_submit=False):
                email = st.text_input("Email")
                password = st.text_input("Password", type="password")
                submitted = st.form_submit_button("Sign In", use_container_width=True)
            if submitted:
                user = authenticate_local_user(email, password)
                if user is None:
                    st.error("We could not sign you in. Check the email, password, and whether this account was created in the current database.")
                    log_api_error("auth", "Failed sign-in attempt", {"email": email.strip().lower()})
                else:
                    user["is_guest"] = False
                    st.session_state.current_user = user
                    log_audit_event(user["email"], user["role"], "sign_in", "user", user["email"], {"source": "sidebar"})
                    st.success("Signed in successfully.")
                    st.rerun()

        with sign_up_tab:
            with st.form("sign_up_form", clear_on_submit=True):
                display_name = st.text_input("Display name")
                email = st.text_input("Email address")
                password = st.text_input("Password", type="password")
                role = st.selectbox("Role", options=["teacher", "student"])
                submitted = st.form_submit_button("Create Profile", use_container_width=True)
            if submitted:
                if not is_valid_email(email):
                    st.error("Enter a valid email address before creating the profile.")
                else:
                    ok, message = create_local_user(email, password, display_name, role)
                    if ok:
                        log_audit_event(email.strip().lower(), role, "sign_up", "user", email.strip().lower(), {"display_name": display_name.strip()})
                        st.success(message)
                    else:
                        st.error(message)

        with guest_tab:
            if st.button("Use Guest Teacher Mode", use_container_width=True):
                st.session_state.current_user = default_guest_user()
                st.rerun()

        if not current_user.get("is_guest"):
            if st.button("Sign Out", use_container_width=True):
                log_audit_event(current_user["email"], current_user["role"], "sign_out", "user", current_user["email"], {})
                st.session_state.current_user = default_guest_user()
                st.session_state.last_attempt = None
                st.rerun()


def render_cloud_status_sidebar() -> None:
    """Render cloud sync status and readiness."""
    status = get_cloud_status()
    with st.sidebar.expander("Cloud Sync", expanded=False):
        if status["enabled"]:
            st.success("Supabase cloud database is configured. The app now uses cloud-first storage.")
        else:
            st.info("Supabase is optional. Configure SUPABASE_URL and SUPABASE_KEY to enable cloud-first storage.")
        st.caption(f"SUPABASE_URL: {'Yes' if status['url_present'] else 'No'}")
        st.caption(f"SUPABASE_KEY: {'Yes' if status['key_present'] else 'No'}")


def render_share_links_sidebar() -> None:
    """Render share-link controls and existing active links."""
    with st.sidebar.expander("Share Links", expanded=False):
        st.session_state.public_app_url = st.text_input(
            "Public app URL",
            value=get_public_app_url(),
            help="Set the deployed app URL that students should open.",
        ).strip() or "http://localhost:8501"
        if "localhost" in st.session_state.public_app_url or "127.0.0.1" in st.session_state.public_app_url:
            st.warning("A localhost URL works only on your computer. Deploy the app first before sending links to students.")

        notice = st.session_state.get("share_notice", "")
        if notice:
            st.caption(notice)

        links = list_share_links(
            limit=20,
            owner_email=get_owner_email(),
            test_uid=get_current_test_uid() or None,
        )
        if not links:
            st.info("No shared links yet for the current test.")
            return

        bulk_targets = st.multiselect(
            "Bulk-select links",
            options=[item["token"] for item in links],
            format_func=lambda token: next((f"{item['variant_name']} | {item['title']}" for item in links if item["token"] == token), token),
        )
        if bulk_targets and st.button("Deactivate selected links", use_container_width=True):
            for token in bulk_targets:
                set_share_link_status(token, False)
            log_event("bulk_deactivate_share_links", "share_link", ",".join(bulk_targets), {"count": len(bulk_targets)})
            st.rerun()

        for item in links:
            with st.container(border=True):
                st.markdown(f"**{item['title']}**")
                st.caption(f"{item['variant_name']} | {'Active' if item['is_active'] else 'Inactive'}")
                if item.get("deadline_at"):
                    st.caption(f"Deadline: {item['deadline_at']}")
                st.caption(
                    f"Max attempts: {'Unlimited' if int(item.get('max_attempts', 1)) == 0 else item.get('max_attempts', 1)}"
                )
                share_payload = load_share_link(item["token"]) or {}
                share_settings = share_payload.get("payload", {}).get("share_settings", {})
                if share_settings:
                    active_rules = []
                    if share_settings.get("require_student_login"):
                        active_rules.append("login required")
                    if share_settings.get("allowed_students"):
                        active_rules.append(f"whitelist {len(share_settings.get('allowed_students', []))}")
                    if share_settings.get("per_student_random_order"):
                        active_rules.append("random order")
                    if int(share_settings.get("timer_minutes", 0) or 0) > 0:
                        active_rules.append(f"{share_settings.get('timer_minutes')} min timer")
                    if share_settings.get("one_question_at_a_time"):
                        active_rules.append("one-question mode")
                    if share_settings.get("block_copy_print"):
                        active_rules.append("copy/print deterrent")
                    if share_settings.get("no_instant_score"):
                        active_rules.append("no instant score")
                    if active_rules:
                        st.caption("Rules: " + " | ".join(active_rules))
                share_url = build_share_url(item["token"])
                st.code(share_url, language=None)
                toggle_label = "Deactivate" if item["is_active"] else "Activate"
                if st.button(toggle_label, key=f"toggle_share_{item['token']}", use_container_width=True):
                    set_share_link_status(item["token"], not bool(item["is_active"]))
                    st.rerun()


def get_public_app_url() -> str:
    """Return the configured base URL for shared links."""
    return st.session_state.get("public_app_url", "http://localhost:8501").rstrip("/")


def build_share_url(token: str) -> str:
    """Build a full shared student URL."""
    return f"{get_public_app_url()}?share={token}"


def maybe_sync_history(record: dict[str, Any]) -> None:
    """Sync a history record to cloud storage if configured."""
    if not is_cloud_enabled():
        return
    return


def maybe_sync_question_bank(record: dict[str, Any]) -> None:
    """Sync a question bank item to cloud storage if configured."""
    if not is_cloud_enabled():
        return
    return


def maybe_sync_attempt(record: dict[str, Any]) -> None:
    """Sync an attempt result to cloud storage in the background if configured."""
    if not is_cloud_enabled():
        return
    return


def get_default_topic(uploaded_name: str | None) -> str:
    """Derive a topic when the user generates from a file only."""
    if not uploaded_name:
        return ""
    stem = os.path.splitext(uploaded_name)[0].replace("_", " ").replace("-", " ").strip()
    return stem.title()


def generate_variant_pack(
    *,
    topic: str,
    question_count: int,
    test_type: str,
    language: str,
    grade_level: str,
    learning_objective: str,
    lesson_stage: str,
    assessment_purpose: str,
    source_material: str,
    source_name: str,
) -> dict[str, dict[str, Any]]:
    """Generate easy, medium, hard, and mixed variants as one pack."""
    sources: dict[str, dict[str, Any]] = {}
    for variant_name, difficulty in (
        ("Variant A", "easy"),
        ("Variant B", "medium"),
        ("Variant C", "hard"),
    ):
        sources[variant_name] = generate_test(
            topic=topic,
            question_count=question_count,
            difficulty=difficulty,
            test_type=test_type,
            language=language,
            grade_level=grade_level,
            learning_objective=learning_objective,
            lesson_stage=lesson_stage,
            assessment_purpose=assessment_purpose,
            source_material=source_material,
            source_name=source_name,
        )
    return build_all_variants(sources)


def render_generator_form() -> tuple[dict[str, Any], bool]:
    """Render all generation controls and return their values."""
    open_section("Create Test")
    top_left, top_middle, top_right = st.columns([2.2, 0.9, 1], gap="large")

    with top_left:
        topic = st.text_input(
            "Topic",
            placeholder="Enter a subject, chapter, or lesson theme",
            help="You can leave this empty if you upload a source file.",
        )

    with top_middle:
        question_count = st.slider(
            "Number of questions",
            min_value=1,
            max_value=10,
            value=5,
        )

    with top_right:
        uploaded_file = st.file_uploader(
            "Source material (optional)",
            type=[extension.lstrip(".") for extension in SUPPORTED_EXTENSIONS],
            help="Upload a PDF, DOCX, or TXT file to generate a test from its contents.",
        )

    with st.expander("Advanced settings", expanded=False):
        row_two_left, row_two_middle, row_two_right = st.columns(3, gap="large")
        row_three_left, row_three_middle, row_three_right = st.columns(3, gap="large")

        with row_two_left:
            grade_level = st.selectbox("Grade level", options=GRADE_LEVEL_OPTIONS)

        with row_two_middle:
            difficulty = st.selectbox(
                "Editor starting difficulty",
                options=list(DIFFICULTY_OPTIONS),
                format_func=lambda key: DIFFICULTY_OPTIONS[key],
                help="All four variants are generated automatically. This setting chooses which difficulty opens first in the editor.",
            )

        with row_two_right:
            test_type = st.selectbox(
                "Test type",
                options=list(TEST_TYPE_OPTIONS),
                format_func=lambda key: TEST_TYPE_OPTIONS[key],
            )

        with row_three_left:
            language = st.selectbox(
                "Language",
                options=list(LANGUAGE_OPTIONS),
                format_func=lambda key: LANGUAGE_OPTIONS[key],
            )

        with row_three_middle:
            lesson_stage = st.selectbox("Lesson stage", options=LESSON_STAGE_OPTIONS)

        with row_three_right:
            assessment_purpose = st.selectbox("Assessment purpose", options=ASSESSMENT_PURPOSE_OPTIONS)

        learning_objective = st.text_area(
            "Learning objective",
            placeholder="Example: Students identify literary devices and justify their answers.",
            height=90,
        )

        subject_tags = st.text_input(
            "Subject tags",
            placeholder="Example: literature, reading, analysis",
            help="Comma-separated tags for the teacher library.",
        )

    st.write("")
    generate_clicked = st.button("Generate Test", use_container_width=True, type="primary")
    close_section()
    return {
        "topic": topic,
        "question_count": question_count,
        "difficulty": difficulty,
        "test_type": test_type,
        "language": language,
        "uploaded_file": uploaded_file,
        "grade_level": grade_level,
        "lesson_stage": lesson_stage,
        "assessment_purpose": assessment_purpose,
        "learning_objective": learning_objective.strip(),
        "subject_tags": format_subject_tags(subject_tags),
    }, generate_clicked


def save_current_test_snapshot(test_data: dict[str, Any], metadata: dict[str, Any]) -> int:
    """Persist the current test snapshot to SQLite."""
    payload = {
        "test_data": test_data,
        "variants": st.session_state.get("generated_variants", {}),
        "metadata": metadata,
    }
    record_id = save_test_record(
        test_uid=metadata.get("test_uid", ""),
        title=test_data.get("title", "Generated Test"),
        topic=metadata.get("topic", ""),
        language=metadata.get("language", "english"),
        difficulty=metadata.get("difficulty", "medium"),
        test_type=metadata.get("test_type", "multiple_choice"),
        grade_level=metadata.get("grade_level", ""),
        assessment_purpose=metadata.get("assessment_purpose", ""),
        owner_email=get_owner_email(),
        source_kind=metadata.get("source_kind", "topic"),
        source_name=metadata.get("source_name", ""),
        subject_tags=format_subject_tags(metadata.get("subject_tags", "")),
        is_favorite=bool(metadata.get("is_favorite", False)),
        payload=payload,
    )
    maybe_sync_history(
        {
            "test_uid": metadata.get("test_uid", ""),
            "title": test_data.get("title", "Generated Test"),
            "topic": metadata.get("topic", ""),
            "language": metadata.get("language", "english"),
            "difficulty": metadata.get("difficulty", "medium"),
            "test_type": metadata.get("test_type", "multiple_choice"),
            "grade_level": metadata.get("grade_level", ""),
            "assessment_purpose": metadata.get("assessment_purpose", ""),
            "source_kind": metadata.get("source_kind", "topic"),
            "source_name": metadata.get("source_name", ""),
            "subject_tags": format_subject_tags(metadata.get("subject_tags", "")),
            "payload": payload,
        }
    )
    return record_id


def extract_source_preview(uploaded_file: Any) -> tuple[str, dict[str, Any]]:
    """Extract and summarize uploaded material."""
    source_material = extract_text_from_uploaded_file(uploaded_file.name, uploaded_file.getvalue())
    preview = source_material[:1200]
    stats = {
        "file_name": uploaded_file.name,
        "character_count": len(source_material),
        "preview_character_count": len(preview),
    }
    return source_material, stats


def handle_generation(form_data: dict[str, Any]) -> None:
    """Validate input, optionally extract uploaded material, and generate a test."""
    allowed, message = can_use_generation()
    if not allowed:
        st.error(message)
        return

    source_material = ""
    source_kind = "topic"
    source_name = ""
    clean_topic = form_data["topic"].strip()

    uploaded_file = form_data["uploaded_file"]
    if uploaded_file is not None:
        try:
            source_material, source_stats = extract_source_preview(uploaded_file)
        except ValueError as error:
            st.error(str(error))
            return

        source_kind = "file"
        source_name = uploaded_file.name
        st.session_state.source_preview = source_material[:1200]
        st.session_state.source_stats = source_stats

        if not clean_topic:
            clean_topic = get_default_topic(uploaded_file.name)
    else:
        st.session_state.source_preview = ""
        st.session_state.source_stats = {}

    if not clean_topic:
        st.error("Please enter a topic or upload a source file before generating a test.")
        return

    try:
        with st.spinner("Generating four classroom variants: easy, medium, hard, and mixed..."):
            generated_variants = generate_variant_pack(
                topic=clean_topic,
                question_count=form_data["question_count"],
                test_type=form_data["test_type"],
                language=form_data["language"],
                grade_level=form_data["grade_level"],
                learning_objective=form_data["learning_objective"],
                lesson_stage=form_data["lesson_stage"],
                assessment_purpose=form_data["assessment_purpose"],
                source_material=source_material,
                source_name=source_name,
            )
    except ValueError as error:
        log_api_error(
            "groq",
            str(error),
            {"topic": clean_topic, "test_type": form_data["test_type"], "language": form_data["language"]},
        )
        st.error(friendly_generation_error_message(error))
        return
    except RuntimeError as error:
        log_api_error(
            "groq",
            str(error),
            {"topic": clean_topic, "test_type": form_data["test_type"], "language": form_data["language"]},
        )
        st.error(friendly_generation_error_message(error))
        return
    except Exception:
        log_api_error(
            "groq",
            "Unexpected generation failure",
            {"topic": clean_topic, "test_type": form_data["test_type"], "language": form_data["language"]},
        )
        st.error("The test could not be generated because of an unexpected platform error. Please retry or check the AI provider logs.")
        return

    test_uid = uuid4().hex
    for variant_name, variant_data in generated_variants.items():
        variant_data["test_uid"] = test_uid
        variant_data["variant_name"] = variant_name

    editor_variant_name = {
        "easy": "Variant A",
        "medium": "Variant B",
        "hard": "Variant C",
    }.get(form_data["difficulty"], "Variant D")
    generated_test = deepcopy(generated_variants[editor_variant_name])
    generated_test["editor_variant_name"] = editor_variant_name
    metadata = {
        "test_uid": test_uid,
        "topic": clean_topic,
        "language": form_data["language"],
        "difficulty": form_data["difficulty"],
        "editor_variant_name": editor_variant_name,
        "test_type": form_data["test_type"],
        "question_count": form_data["question_count"],
        "source_kind": source_kind,
        "source_name": source_name,
        "grade_level": form_data["grade_level"],
        "learning_objective": form_data["learning_objective"],
        "lesson_stage": form_data["lesson_stage"],
        "assessment_purpose": form_data["assessment_purpose"],
        "subject_tags": form_data.get("subject_tags", ""),
        "is_favorite": False,
    }

    st.session_state.generated_test = generated_test
    st.session_state.generated_variants = generated_variants
    st.session_state.generated_topic = clean_topic
    st.session_state.test_metadata = metadata
    st.session_state.editor_version += 1
    st.session_state.quality_report = analyze_test_quality(
        generated_test,
        expected_count=form_data["question_count"],
    )
    st.session_state.last_attempt = None
    st.session_state.active_flow_step = "Review"
    fallback_used = any(bool(item.get("fallback_mode")) for item in generated_variants.values())

    record_id = save_current_test_snapshot(generated_test, metadata)
    record_usage_event(
        get_owner_email(),
        "generation",
        1,
        {"topic": clean_topic, "test_uid": test_uid, "test_type": form_data["test_type"]},
    )
    log_event("generate_test", "test", test_uid, {"topic": clean_topic, "record_id": record_id})
    st.session_state.history_notice = f"Saved to local history as record #{record_id}."
    st.session_state.last_autosave_signature = build_payload_signature(generated_test, metadata)
    if fallback_used:
        st.warning("The cloud AI service was unavailable, so a local fallback generator created the test pack. Review the questions carefully before sharing.")
    else:
        st.success("Four variants were generated successfully. You can now review Variant D, export all variants, and track analytics per test.")


def build_file_base_name(topic: str, test_type: str, language: str, export_mode: str, variant_name: str) -> str:
    """Build a safe file base name for downloaded files."""
    topic_part = "".join(char if char.isalnum() else "_" for char in topic.strip().lower())
    topic_part = "_".join(filter(None, topic_part.split("_"))) or "teacher_test"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    variant_part = variant_name.lower().replace(" ", "_")
    return f"{topic_part}_{test_type}_{language}_{variant_part}_{export_mode}_{timestamp}"


def get_editor_prefix() -> str:
    """Return a stable widget prefix for the current generated test."""
    return f"editor_{st.session_state.editor_version}"


def build_context_lines(test_data: dict[str, Any], variant_name: str = "") -> list[str]:
    """Build context lines that can be included in exports."""
    pairs = [
        ("Topic", test_data.get("topic", "")),
        ("Grade level", test_data.get("grade_level", "")),
        ("Learning objective", test_data.get("learning_objective", "")),
        ("Lesson stage", test_data.get("lesson_stage", "")),
        ("Assessment purpose", test_data.get("assessment_purpose", "")),
        ("Variant", variant_name or test_data.get("variant_name", "")),
        ("Variant difficulty", test_data.get("variant_label", test_data.get("variant_difficulty", ""))),
    ]
    return [f"{label}: {value}" for label, value in pairs if str(value).strip()]


def format_correct_answer(question: dict[str, Any]) -> str:
    """Format the correct answer for export."""
    question_type = question["type"]
    if question_type in {"multiple_choice", "true_false"}:
        options = question.get("options", [])
        answer = question.get("correct_answer", "")
        for index, option in enumerate(options):
            if option == answer:
                return f"{chr(65 + index)}) {option}"
        return answer

    if question_type == "matching":
        pairs = question.get("pairs", [])
        if not pairs:
            return ""
        return "; ".join(f"{pair['left']} -> {pair['right']}" for pair in pairs)

    return question.get("correct_answer", "")


def build_matching_student_columns(question: dict[str, Any]) -> tuple[list[str], list[str]]:
    """Build separate columns for student-facing matching tasks."""
    pairs = question.get("pairs", [])
    left_items = [pair["left"] for pair in pairs]
    right_items = [pair["right"] for pair in pairs]
    if len(right_items) > 1:
        right_items = right_items[1:] + right_items[:1]
    return left_items, right_items


def build_text_export(test_data: dict[str, Any], include_answers: bool, variant_name: str) -> str:
    """Create a plain text export for the current test."""
    lines = [test_data["title"], ""]
    context_lines = build_context_lines(test_data, variant_name)
    if context_lines:
        lines.extend(context_lines)
        lines.append("")

    instructions = test_data.get("instructions", "").strip()
    if instructions:
        lines.extend([instructions, ""])

    for index, question in enumerate(test_data["questions"], start=1):
        lines.append(f"Question {index}: {question['question']}")
        if question["type"] in {"multiple_choice", "true_false"}:
            for option_index, option in enumerate(question.get("options", [])):
                lines.append(f"{chr(65 + option_index)}) {option}")
            if include_answers:
                lines.append(f"Correct answer: {format_correct_answer(question)}")
                lines.append(f"Explanation: {question.get('explanation', '')}")
        elif question["type"] == "short_answer":
            if include_answers:
                lines.append(f"Correct answer: {question.get('correct_answer', '')}")
                lines.append(f"Explanation: {question.get('explanation', '')}")
            else:
                lines.append("Answer: ____________________")
        elif question["type"] == "matching":
            if include_answers:
                lines.append("Pairs:")
                for pair_index, pair in enumerate(question.get("pairs", []), start=1):
                    lines.append(f"{pair_index}. {pair['left']} -> {pair['right']}")
                lines.append(f"Correct answer: {format_correct_answer(question)}")
                lines.append(f"Explanation: {question.get('explanation', '')}")
            else:
                left_items, right_items = build_matching_student_columns(question)
                lines.append("Column A:")
                for pair_index, value in enumerate(left_items, start=1):
                    lines.append(f"{pair_index}. {value}")
                lines.append("Column B:")
                for pair_index, value in enumerate(right_items, start=1):
                    lines.append(f"{chr(64 + pair_index)}) {value}")
        lines.append("")

    return "\n".join(lines).strip()


def build_docx_export(test_data: dict[str, Any], include_answers: bool, variant_name: str) -> bytes:
    """Create a DOCX document in memory."""
    document = Document()
    title = document.add_heading(test_data["title"], level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for line in build_context_lines(test_data, variant_name):
        document.add_paragraph(line)

    instructions = test_data.get("instructions", "").strip()
    if instructions:
        document.add_paragraph(instructions)

    for index, question in enumerate(test_data["questions"], start=1):
        document.add_heading(f"Question {index}", level=2)
        document.add_paragraph(question["question"])
        if question["type"] in {"multiple_choice", "true_false"}:
            for option_index, option in enumerate(question.get("options", [])):
                document.add_paragraph(f"{chr(65 + option_index)}) {option}")
            if include_answers:
                document.add_paragraph(f"Correct answer: {format_correct_answer(question)}")
                document.add_paragraph(f"Explanation: {question.get('explanation', '')}")
        elif question["type"] == "short_answer":
            if include_answers:
                document.add_paragraph(f"Correct answer: {question.get('correct_answer', '')}")
                document.add_paragraph(f"Explanation: {question.get('explanation', '')}")
            else:
                document.add_paragraph("Answer: ____________________")
        elif question["type"] == "matching":
            if include_answers:
                for pair_index, pair in enumerate(question.get("pairs", []), start=1):
                    document.add_paragraph(f"{pair_index}. {pair['left']} -> {pair['right']}")
                document.add_paragraph(f"Correct answer: {format_correct_answer(question)}")
                document.add_paragraph(f"Explanation: {question.get('explanation', '')}")
            else:
                left_items, right_items = build_matching_student_columns(question)
                document.add_paragraph("Column A:")
                for pair_index, value in enumerate(left_items, start=1):
                    document.add_paragraph(f"{pair_index}. {value}")
                document.add_paragraph("Column B:")
                for pair_index, value in enumerate(right_items, start=1):
                    document.add_paragraph(f"{chr(64 + pair_index)}) {value}")

    for section in document.sections:
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(18)
        section.right_margin = Mm(18)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@lru_cache(maxsize=1)
def get_pdf_font_name() -> str:
    """Register a Unicode-capable font for PDF export when available."""
    font_candidates = [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
    ]
    for path in font_candidates:
        if os.path.exists(path):
            pdfmetrics.registerFont(TTFont("TeacherGeneratorFont", path))
            return "TeacherGeneratorFont"
    return "Helvetica"


def build_pdf_export(test_data: dict[str, Any], include_answers: bool, variant_name: str) -> bytes:
    """Create a PDF document in memory."""
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    base_font = get_pdf_font_name()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name="TeacherTitle",
        parent=styles["Title"],
        fontName=base_font,
        fontSize=18,
        leading=22,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        name="TeacherBody",
        parent=styles["BodyText"],
        fontName=base_font,
        fontSize=11,
        leading=15,
        spaceAfter=6,
    )
    question_style = ParagraphStyle(
        name="TeacherQuestion",
        parent=body_style,
        fontName=base_font,
        fontSize=12,
        leading=16,
        spaceBefore=8,
        spaceAfter=6,
    )

    story: list[Any] = [Paragraph(escape(test_data["title"]), title_style)]
    for line in build_context_lines(test_data, variant_name):
        story.append(Paragraph(escape(line), body_style))

    instructions = test_data.get("instructions", "").strip()
    if instructions:
        story.append(Paragraph(escape(instructions).replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 4))

    for index, question in enumerate(test_data["questions"], start=1):
        story.append(Paragraph(escape(f"Question {index}: {question['question']}"), question_style))
        if question["type"] in {"multiple_choice", "true_false"}:
            for option_index, option in enumerate(question.get("options", [])):
                story.append(Paragraph(escape(f"{chr(65 + option_index)}) {option}"), body_style))
            if include_answers:
                story.append(Paragraph(escape(f"Correct answer: {format_correct_answer(question)}"), body_style))
                story.append(Paragraph(escape(f"Explanation: {question.get('explanation', '')}"), body_style))
        elif question["type"] == "short_answer":
            answer_line = f"Correct answer: {question.get('correct_answer', '')}" if include_answers else "Answer: ____________________"
            story.append(Paragraph(escape(answer_line), body_style))
            if include_answers:
                story.append(Paragraph(escape(f"Explanation: {question.get('explanation', '')}"), body_style))
        elif question["type"] == "matching":
            if include_answers:
                for pair_index, pair in enumerate(question.get("pairs", []), start=1):
                    story.append(Paragraph(escape(f"{pair_index}. {pair['left']} -> {pair['right']}"), body_style))
                story.append(Paragraph(escape(f"Correct answer: {format_correct_answer(question)}"), body_style))
                story.append(Paragraph(escape(f"Explanation: {question.get('explanation', '')}"), body_style))
            else:
                left_items, right_items = build_matching_student_columns(question)
                story.append(Paragraph("Column A:", body_style))
                for pair_index, value in enumerate(left_items, start=1):
                    story.append(Paragraph(escape(f"{pair_index}. {value}"), body_style))
                story.append(Paragraph("Column B:", body_style))
                for pair_index, value in enumerate(right_items, start=1):
                    story.append(Paragraph(escape(f"{chr(64 + pair_index)}) {value}"), body_style))
        story.append(Spacer(1, 6))

    document.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def render_source_preview(test_data: dict[str, Any]) -> None:
    """Render a preview of the uploaded material or source summary."""
    source_preview = st.session_state.get("source_preview", "")
    source_stats = st.session_state.get("source_stats", {})
    source_summary = test_data.get("source_summary", "")
    key_concepts = test_data.get("key_concepts", [])
    if not source_preview and not source_summary:
        return

    with st.expander("Source Material Intelligence", expanded=False):
        if source_stats:
            st.caption(
                f"Source: {source_stats.get('file_name', '')} | "
                f"Characters extracted: {source_stats.get('character_count', 0)}"
            )
        if source_summary:
            st.markdown("**Source summary used for generation**")
            st.write(source_summary[:1500])
        if key_concepts:
            st.markdown("**Key concepts**")
            st.write(", ".join(key_concepts))
        if source_preview:
            st.text_area(
                "Extracted preview",
                value=source_preview,
                height=180,
                disabled=True,
                label_visibility="collapsed",
            )


def render_question_editor(question: dict[str, Any], index: int) -> dict[str, Any]:
    """Render editable controls for a single question and return the updated data."""
    prefix = get_editor_prefix()
    question_type = question["type"]
    type_label = TEST_TYPE_OPTIONS.get(question_type, question_type.replace("_", " ").title())

    with st.container(border=True):
        st.markdown(f"**Question {index + 1}**")
        st.caption(type_label)

        question_text = st.text_area(
            "Question text",
            value=question.get("question", ""),
            key=f"{prefix}_question_{index}",
            height=100,
        ).strip()

        skill_tag = st.text_input(
            "Skill tag",
            value=question.get("skill_tag", ""),
            key=f"{prefix}_question_{index}_skill",
        ).strip()

        explanation = st.text_area(
            "Teacher explanation",
            value=question.get("explanation", ""),
            key=f"{prefix}_question_{index}_explanation",
            height=90,
        ).strip()

        if question_type == "multiple_choice":
            options: list[str] = []
            default_values = question.get("options", ["", "", "", ""])
            for option_index in range(4):
                option_value = st.text_input(
                    f"Option {chr(65 + option_index)}",
                    value=default_values[option_index] if option_index < len(default_values) else "",
                    key=f"{prefix}_question_{index}_option_{option_index}",
                ).strip()
                options.append(option_value)

            answer_choices = [option for option in options if option] or [""]
            current_answer = question.get("correct_answer", "")
            if current_answer and current_answer not in answer_choices:
                answer_choices.append(current_answer)

            selected_answer = st.selectbox(
                "Correct answer",
                options=answer_choices,
                index=answer_choices.index(current_answer) if current_answer in answer_choices else 0,
                key=f"{prefix}_question_{index}_answer",
            )

            return {
                "type": question_type,
                "question": question_text,
                "options": options,
                "correct_answer": selected_answer,
                "skill_tag": skill_tag,
                "explanation": explanation,
                "pairs": [],
            }

        if question_type == "true_false":
            default_options = question.get("options", [])
            if len(default_options) < 2:
                default_options = ["True", "False"]

            options = []
            option_cols = st.columns(2, gap="large")
            for option_index in range(2):
                with option_cols[option_index]:
                    option_value = st.text_input(
                        f"Option {chr(65 + option_index)}",
                        value=default_options[option_index],
                        key=f"{prefix}_question_{index}_tf_option_{option_index}",
                    ).strip()
                    options.append(option_value)

            answer_choices = [option for option in options if option] or [""]
            current_answer = question.get("correct_answer", "")
            if current_answer and current_answer not in answer_choices:
                answer_choices.append(current_answer)

            selected_answer = st.selectbox(
                "Correct answer",
                options=answer_choices,
                index=answer_choices.index(current_answer) if current_answer in answer_choices else 0,
                key=f"{prefix}_question_{index}_tf_answer",
            )

            return {
                "type": question_type,
                "question": question_text,
                "options": options,
                "correct_answer": selected_answer,
                "skill_tag": skill_tag,
                "explanation": explanation,
                "pairs": [],
            }

        if question_type == "short_answer":
            correct_answer = st.text_input(
                "Correct answer",
                value=question.get("correct_answer", ""),
                key=f"{prefix}_question_{index}_short_answer",
            ).strip()

            return {
                "type": question_type,
                "question": question_text,
                "options": [],
                "correct_answer": correct_answer,
                "skill_tag": skill_tag,
                "explanation": explanation,
                "pairs": [],
            }

        pair_values = question.get("pairs", [])
        pair_count = st.number_input(
            "Number of pairs",
            min_value=2,
            max_value=8,
            value=max(2, len(pair_values)),
            step=1,
            key=f"{prefix}_question_{index}_pair_count",
        )

        edited_pairs = []
        for pair_index in range(int(pair_count)):
            left_default = pair_values[pair_index]["left"] if pair_index < len(pair_values) else ""
            right_default = pair_values[pair_index]["right"] if pair_index < len(pair_values) else ""
            left_col, right_col = st.columns(2, gap="large")
            with left_col:
                left_value = st.text_input(
                    f"Left item {pair_index + 1}",
                    value=left_default,
                    key=f"{prefix}_question_{index}_pair_left_{pair_index}",
                ).strip()
            with right_col:
                right_value = st.text_input(
                    f"Right item {pair_index + 1}",
                    value=right_default,
                    key=f"{prefix}_question_{index}_pair_right_{pair_index}",
                ).strip()
            edited_pairs.append({"left": left_value, "right": right_value})

        return {
            "type": question_type,
            "question": question_text,
            "options": [],
            "correct_answer": "",
            "skill_tag": skill_tag,
            "explanation": explanation,
            "pairs": edited_pairs,
        }


def save_question_to_bank(question: dict[str, Any]) -> None:
    """Save a question into the question bank."""
    record_id = save_question_bank_item(
        question_text=question.get("question", ""),
        question_type=question.get("type", ""),
        topic=st.session_state.test_metadata.get("topic", ""),
        skill_tag=question.get("skill_tag", ""),
        owner_email=get_owner_email(),
        payload=question,
    )
    st.session_state.question_bank_notice = f"Saved question to bank as item #{record_id}."
    maybe_sync_question_bank(
        {
            "question_text": question.get("question", ""),
            "question_type": question.get("type", ""),
            "topic": st.session_state.test_metadata.get("topic", ""),
            "skill_tag": question.get("skill_tag", ""),
            "owner_email": get_owner_email(),
            "payload": question,
        }
    )
    st.success(f"Question saved to bank as item #{record_id}.")


def regenerate_single_question(question_index: int) -> None:
    """Regenerate one question using the current test metadata."""
    if st.session_state.generated_test is None:
        return
    metadata = st.session_state.test_metadata
    current_questions = deepcopy(st.session_state.generated_test.get("questions", []))
    if question_index >= len(current_questions):
        return

    try:
        replacement_test = generate_test(
            topic=metadata.get("topic", ""),
            question_count=1,
            difficulty=metadata.get("difficulty", "medium"),
            test_type=metadata.get("test_type", "multiple_choice"),
            language=metadata.get("language", "english"),
            grade_level=metadata.get("grade_level", ""),
            learning_objective=metadata.get("learning_objective", ""),
            lesson_stage=metadata.get("lesson_stage", ""),
            assessment_purpose=metadata.get("assessment_purpose", ""),
            source_material=st.session_state.get("source_preview", ""),
            source_name=metadata.get("source_name", ""),
        )
    except Exception as error:
        st.error(f"Question regeneration failed: {error}")
        return

    replacement_question = replacement_test.get("questions", [{}])[0]
    current_questions[question_index] = replacement_question
    st.session_state.generated_test["questions"] = current_questions
    st.session_state.editor_version += 1
    st.session_state.quality_report = analyze_test_quality(
        st.session_state.generated_test,
        expected_count=metadata.get("question_count"),
    )
    st.success(f"Question {question_index + 1} was regenerated.")
    st.rerun()


def render_test_editor(test_data: dict[str, Any]) -> dict[str, Any]:
    """Render the editable generated test and return the updated structure."""
    prefix = get_editor_prefix()
    editable_test = deepcopy(test_data)

    open_section("Teacher Workspace")

    title = st.text_input(
        "Test title",
        value=test_data.get("title", "Generated Test"),
        key=f"{prefix}_title",
    ).strip()
    instructions = st.text_area(
        "Instructions",
        value=test_data.get("instructions", ""),
        key=f"{prefix}_instructions",
        height=90,
    ).strip()

    editable_test["title"] = title or "Generated Test"
    editable_test["instructions"] = instructions

    edited_questions = []
    for index, question in enumerate(test_data.get("questions", [])):
        edited_question = render_question_editor(question, index)
        action_col1, action_col2, action_col3 = st.columns([1, 1, 4])
        with action_col1:
            if st.button("Save to Bank", key=f"{prefix}_bank_{index}", use_container_width=True):
                save_question_to_bank(edited_question)
        with action_col2:
            if st.button("Regenerate", key=f"{prefix}_regen_{index}", use_container_width=True):
                regenerate_single_question(index)
        edited_questions.append(edited_question)

    editable_test["questions"] = edited_questions
    close_section()
    return editable_test


def render_metadata_summary() -> None:
    """Show metadata for the current test."""
    metadata = st.session_state.test_metadata
    if not metadata:
        return

    open_section("Test Snapshot")
    info_cols = st.columns(6, gap="small")
    info_cols[0].metric("Topic", metadata.get("topic", ""))
    info_cols[1].metric("Grade", metadata.get("grade_level", ""))
    info_cols[2].metric("Difficulty", DIFFICULTY_OPTIONS.get(metadata.get("difficulty", ""), ""))
    info_cols[3].metric("Type", TEST_TYPE_OPTIONS.get(metadata.get("test_type", ""), ""))
    info_cols[4].metric("Language", LANGUAGE_OPTIONS.get(metadata.get("language", ""), ""))
    info_cols[5].metric("Source", metadata.get("source_name") or metadata.get("source_kind", "topic").title())
    if metadata.get("learning_objective"):
        st.caption(f"Learning objective: {metadata['learning_objective']}")
    if metadata.get("subject_tags"):
        st.caption(f"Subject tags: {metadata['subject_tags']}")
    st.caption(
        f"Lesson stage: {metadata.get('lesson_stage', '')} | "
        f"Assessment purpose: {metadata.get('assessment_purpose', '')}"
    )
    close_section()


def render_quality_report(quality_report: dict[str, Any]) -> None:
    """Render a quality summary for the current test."""
    open_section("Quality Report")
    top_left, top_middle, top_right = st.columns(3, gap="large")
    top_left.metric("Quality score", f"{quality_report['score']}/100")
    top_middle.metric("Warnings", len(quality_report["warnings"]))
    top_right.metric("Blocking issues", len(quality_report["blocking_issues"]))

    if quality_report["is_export_ready"]:
        st.success("The current test structure is ready for export and student practice.")
    else:
        st.warning("Fix the blocking issues below before export or student practice.")

    issue_col, warning_col, strength_col = st.columns(3, gap="large")
    with issue_col:
        st.markdown("**Blocking issues**")
        if quality_report["blocking_issues"]:
            for item in quality_report["blocking_issues"]:
                st.error(item)
        else:
            st.info("No blocking issues.")
    with warning_col:
        st.markdown("**Warnings**")
        if quality_report["warnings"]:
            for item in quality_report["warnings"]:
                st.warning(item)
        else:
            st.info("No warnings.")
    with strength_col:
        st.markdown("**Strengths**")
        if quality_report["strengths"]:
            for item in quality_report["strengths"]:
                st.success(item)
        else:
            st.info("No strengths recorded yet.")
    close_section()


def parse_deadline(deadline_value: str) -> datetime | None:
    """Parse a share-link deadline from an ISO-like datetime-local string."""
    if not str(deadline_value).strip():
        return None
    try:
        return datetime.fromisoformat(str(deadline_value).strip())
    except ValueError:
        return None


def render_test_status_banner() -> None:
    """Render a small workspace status bar."""
    if not st.session_state.test_metadata:
        return
    status_col1, status_col2, status_col3 = st.columns([1, 1, 2], gap="large")
    status_col1.metric("Status", get_status_label())
    status_col2.metric("Flow", "Review → Share → Analyze")
    status_col3.caption(
        f"Test UID: {get_current_test_uid() or 'not created'} | "
        f"Last autosave: {'Ready' if st.session_state.get('last_autosave_signature') else 'Pending'}"
    )


def render_quick_actions() -> None:
    """Render compact teacher actions in a More menu."""
    with st.popover("More"):
        current_favorite = bool(st.session_state.test_metadata.get("is_favorite", False))
        favorite_label = "Remove from favorites" if current_favorite else "Add to favorites"
        if st.button(favorite_label, use_container_width=True):
            current_uid = get_current_test_uid()
            if current_uid:
                new_value = not current_favorite
                set_test_favorite(current_uid, get_owner_email(), new_value)
                st.session_state.test_metadata["is_favorite"] = new_value
                st.rerun()
        if st.button("Duplicate test", use_container_width=True):
            duplicate_current_test()
            record_id = save_current_test_snapshot(st.session_state.generated_test, st.session_state.test_metadata)
            st.session_state.history_notice = f"Duplicated as record #{record_id}."
            st.rerun()
        if st.button("Archive test", use_container_width=True):
            archive_current_test()
            st.rerun()
        if st.session_state.get("last_created_share_url"):
            st.caption("Latest share link")
            st.code(st.session_state["last_created_share_url"], language=None)
        if st.button("Close workspace", use_container_width=True):
            clear_workspace()
            st.rerun()


def render_variant_export_block(variant_name: str, variant_data: dict[str, Any], disable: bool) -> None:
    """Render exports for one variant."""
    st.markdown(f"**{variant_name}**")
    share_col, note_col = st.columns([1, 2], gap="large")
    with share_col:
        if disable:
            st.caption("Share becomes available after the test passes quality checks.")
        else:
            with st.popover(f"Share {variant_name}"):
                max_attempts = st.number_input(
                    "Max attempts per student",
                    min_value=0,
                    max_value=10,
                    value=1,
                    step=1,
                    key=f"share_limit_{variant_name}",
                    help="Use 0 for unlimited attempts.",
                )
                deadline_value = st.text_input(
                    "Deadline (optional)",
                    value="",
                    key=f"share_deadline_{variant_name}",
                    placeholder="2026-04-30T18:00",
                    help="Use local datetime format YYYY-MM-DDTHH:MM.",
                ).strip()
                require_student_login = st.checkbox(
                    "Require student sign-in",
                    value=True,
                    key=f"share_require_login_{variant_name}",
                    help="Only authenticated student accounts can open and submit this link.",
                )
                if require_student_login:
                    st.caption("Authenticated student links are limited to one attempt per student account.")
                whitelist_raw = st.text_area(
                    "Allowed student emails (optional)",
                    value="",
                    key=f"share_whitelist_{variant_name}",
                    placeholder="student1@example.com\nstudent2@example.com",
                    help="Only these student accounts can open the test when sign-in is required.",
                    height=90,
                )
                group_items = list_student_groups(get_owner_email())
                group_map = {item["name"]: int(item["id"]) for item in group_items}
                selected_group_name = st.selectbox(
                    "Or use imported group",
                    options=["None"] + list(group_map),
                    key=f"share_group_{variant_name}",
                    help="Imported student emails from this group will be added to the whitelist automatically.",
                )
                per_student_random_order = st.checkbox(
                    "Randomize order for each student",
                    value=True,
                    key=f"share_random_order_{variant_name}",
                    help="Each student gets the same content, but in a different deterministic order.",
                )
                timer_minutes = st.number_input(
                    "Timer in minutes (0 = no timer)",
                    min_value=0,
                    max_value=240,
                    value=20,
                    step=5,
                    key=f"share_timer_{variant_name}",
                )
                one_question_at_a_time = st.checkbox(
                    "Show one question at a time",
                    value=True,
                    key=f"share_one_question_{variant_name}",
                )
                block_copy_print = st.checkbox(
                    "Soft block copy / print",
                    value=True,
                    key=f"share_block_copy_{variant_name}",
                    help="This is a soft deterrent, not a full security barrier.",
                )
                reveal_score_after_submit = st.checkbox(
                    "Show score after submit",
                    value=True,
                    key=f"share_show_score_{variant_name}",
                    help="Students will see only their score summary, not correct answers.",
                )
                no_instant_score = st.checkbox(
                    "No instant score",
                    value=False,
                    key=f"share_no_score_{variant_name}",
                    help="Students will only see a submission confirmation after finishing the test.",
                )
                if st.button("Create share link", key=f"share_create_{variant_name}", use_container_width=True):
                    effective_max_attempts = 1 if require_student_login else int(max_attempts)
                    whitelist_students = parse_whitelist(whitelist_raw)
                    if selected_group_name != "None":
                        for roster_row in list_group_students(get_owner_email(), group_map[selected_group_name]):
                            roster_email = str(roster_row.get("email", "")).strip().lower()
                            if roster_email and roster_email not in whitelist_students:
                                whitelist_students.append(roster_email)
                    token = create_share_link(
                        test_uid=get_current_test_uid(),
                        title=variant_data.get("title", "Shared Test"),
                        variant_name=variant_name,
                        owner_email=get_owner_email(),
                        payload={
                            "variant_data": variant_data,
                            "share_settings": {
                                "require_student_login": require_student_login,
                                "reveal_score_after_submit": False if no_instant_score else reveal_score_after_submit,
                                "allowed_students": whitelist_students,
                                "per_student_random_order": per_student_random_order,
                                "timer_minutes": int(timer_minutes),
                                "one_question_at_a_time": one_question_at_a_time,
                                "block_copy_print": block_copy_print,
                                "no_instant_score": no_instant_score,
                            },
                        },
                        max_attempts=effective_max_attempts,
                        deadline_at=deadline_value,
                    )
                    share_url = build_share_url(token)
                    copy_share_link_value(share_url)
                    log_event("create_share_link", "share_link", token, {"variant_name": variant_name, "max_attempts": effective_max_attempts})
                    record_usage_event(get_owner_email(), "share_link", 1, {"variant_name": variant_name})
                    st.success("Share link created.")
                    st.code(share_url, language=None)
    with note_col:
        st.caption("Share links open the student page directly. Export actions are kept below in separate teacher and student views.")
    student_tab, teacher_tab = st.tabs(["Student Version", "Teacher Version"])
    for tab_name, include_answers in ((student_tab, False), (teacher_tab, True)):
        with tab_name:
            file_base = build_file_base_name(
                topic=st.session_state.generated_topic,
                test_type=variant_data.get("test_type", "test"),
                language=variant_data.get("language", "english"),
                export_mode="teacher" if include_answers else "student",
                variant_name=variant_name,
            )
            txt_bytes = build_text_export(variant_data, include_answers=include_answers, variant_name=variant_name).encode("utf-8")
            pdf_bytes = build_pdf_export(variant_data, include_answers=include_answers, variant_name=variant_name)
            docx_bytes = build_docx_export(variant_data, include_answers=include_answers, variant_name=variant_name)

            col1, col2, col3 = st.columns(3, gap="large")
            with col1:
                st.download_button("Export TXT", txt_bytes, f"{file_base}.txt", "text/plain", use_container_width=True, disabled=disable, key=f"{variant_name}_{include_answers}_txt")
            with col2:
                st.download_button("Export PDF", pdf_bytes, f"{file_base}.pdf", "application/pdf", use_container_width=True, disabled=disable, key=f"{variant_name}_{include_answers}_pdf")
            with col3:
                st.download_button("Export DOCX", docx_bytes, f"{file_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, disabled=disable, key=f"{variant_name}_{include_answers}_docx")


def get_effective_variants(edited_test: dict[str, Any]) -> dict[str, dict[str, Any]]:
    """Return the active four-variant pack for the current test."""
    variants = deepcopy(st.session_state.get("generated_variants", {}))
    if not variants:
        fallback_sources = {
            "Variant A": deepcopy(edited_test),
            "Variant B": deepcopy(edited_test),
            "Variant C": deepcopy(edited_test),
        }
        variants = build_all_variants(fallback_sources)

    editor_variant_name = st.session_state.test_metadata.get("editor_variant_name", "Variant D")
    variants[editor_variant_name] = deepcopy(edited_test)
    variants[editor_variant_name]["variant_name"] = editor_variant_name
    if all(name in variants for name in ("Variant A", "Variant B", "Variant C")):
        sources = {
            "Variant A": variants["Variant A"],
            "Variant B": variants["Variant B"],
            "Variant C": variants["Variant C"],
        }
        rebuilt = build_all_variants(sources)
        if editor_variant_name == "Variant D":
            rebuilt["Variant D"] = deepcopy(edited_test)
        variants = rebuilt
    st.session_state.generated_variants = variants
    return variants


def render_variants_section(test_data: dict[str, Any], disable: bool) -> dict[str, dict[str, Any]]:
    """Render the four generated variants and export blocks."""
    open_section("Classroom Variants")
    st.subheader("Variants A / B / C / D")
    st.caption("Variant A is easy, Variant B is medium, Variant C is hard, and Variant D is a mixed classroom version.")
    variants = get_effective_variants(test_data)
    tabs = st.tabs(list(variants))
    for tab, (variant_name, variant_data) in zip(tabs, variants.items()):
        with tab:
            st.caption(
                f"{variant_data.get('variant_label', variant_name)} variant. "
                "Questions and options are arranged for classroom use."
            )
            render_variant_export_block(variant_name, variant_data, disable=disable)
    close_section()
    return variants


def render_save_snapshot_button(test_data: dict[str, Any]) -> None:
    """Render a button to save the current edited version to history."""
    if st.button("Save Current Version to History", use_container_width=True):
        record_id = save_current_test_snapshot(test_data, st.session_state.test_metadata)
        st.session_state.history_notice = f"Saved to local history as record #{record_id}."
        st.success(f"Current version saved to local history as record #{record_id}.")


def get_student_widget_prefix(variant_name: str, index: int) -> str:
    """Return the widget prefix used by student-mode questions."""
    return f"student_{variant_name}_{index}"


def is_question_answered(question: dict[str, Any], index: int, variant_name: str) -> bool:
    """Return whether the current student has answered this question."""
    key_prefix = get_student_widget_prefix(variant_name, index)
    if question["type"] in {"multiple_choice", "true_false", "short_answer"}:
        return bool(str(st.session_state.get(f"{key_prefix}_answer", "")).strip())
    if question["type"] == "matching":
        for pair in question.get("pairs", []):
            if not str(st.session_state.get(f"{key_prefix}_match_{pair['left']}", "")).strip():
                return False
        return True
    return False


def count_completed_answers(variant_data: dict[str, Any], variant_name: str) -> int:
    """Count how many questions the current student has answered."""
    return sum(
        1
        for index, question in enumerate(variant_data.get("questions", []))
        if is_question_answered(question, index, variant_name)
    )


def build_submission_key(share_token: str, student_name: str, responses: dict[str, Any]) -> str:
    """Build an idempotency key for one student submission."""
    raw = json.dumps(
        {
            "share_token": share_token,
            "student_name": student_name.strip().lower(),
            "responses": responses,
        },
        ensure_ascii=False,
        sort_keys=True,
    )
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def parse_whitelist(raw_value: str) -> list[str]:
    """Parse one whitelist textarea into lowercase student identifiers."""
    lines = []
    for line in str(raw_value).replace(",", "\n").splitlines():
        cleaned = line.strip().lower()
        if cleaned and cleaned not in lines:
            lines.append(cleaned)
    return lines


def build_personalized_variant(
    variant_data: dict[str, Any],
    share_token: str,
    student_identity: str,
    enable_random_order: bool,
) -> dict[str, Any]:
    """Return a deterministic student-specific variant order."""
    personalized = deepcopy(variant_data)
    if not enable_random_order:
        return personalized

    seed_source = f"{share_token}:{student_identity.lower()}"
    rng = random.Random(seed_source)
    questions = deepcopy(personalized.get("questions", []))
    rng.shuffle(questions)
    updated_questions = []
    for question in questions:
        updated_question = deepcopy(question)
        if updated_question.get("type") in {"multiple_choice", "true_false"}:
            options = list(updated_question.get("options", []))
            rng.shuffle(options)
            updated_question["options"] = options
        elif updated_question.get("type") == "matching":
            pairs = list(updated_question.get("pairs", []))
            rng.shuffle(pairs)
            updated_question["pairs"] = pairs
        updated_questions.append(updated_question)
    personalized["questions"] = updated_questions
    return personalized


def render_soft_exam_protection(watermark_text: str, block_copy_print: bool) -> None:
    """Inject lightweight anti-cheat UI protections for the student page."""
    safe_watermark = escape(watermark_text)
    script = ""
    if block_copy_print:
        script = """
        <script>
        document.addEventListener("contextmenu", function(event) { event.preventDefault(); });
        document.addEventListener("copy", function(event) { event.preventDefault(); });
        document.addEventListener("cut", function(event) { event.preventDefault(); });
        document.addEventListener("keydown", function(event) {
          const key = (event.key || "").toLowerCase();
          if ((event.ctrlKey || event.metaKey) && ["c", "p", "s", "u"].includes(key)) {
            event.preventDefault();
          }
          if (event.key === "PrintScreen") {
            event.preventDefault();
          }
        });
        </script>
        """
    components.html(
        f"""
        <style>
        .student-watermark {{
            position: fixed;
            inset: 0;
            pointer-events: none;
            display: flex;
            align-items: center;
            justify-content: center;
            opacity: 0.08;
            font-size: 3.2rem;
            letter-spacing: 0.08em;
            transform: rotate(-22deg);
            color: #f2d7cf;
            z-index: 0;
            text-transform: uppercase;
            text-align: center;
            white-space: pre-wrap;
        }}
        @media (max-width: 900px) {{
            .student-watermark {{
                font-size: 2rem;
            }}
        }}
        </style>
        <div class="student-watermark">{safe_watermark}</div>
        {script}
        """,
        height=0,
        width=0,
    )


def get_exam_timer_state(share_token: str, student_key: str, minutes_limit: int) -> tuple[datetime | None, int]:
    """Return exam start and seconds left for the authenticated student."""
    if minutes_limit <= 0 or not share_token.strip() or not student_key.strip():
        return None, 0
    state_key = f"exam_started_{share_token}_{student_key.lower()}"
    if state_key not in st.session_state:
        st.session_state[state_key] = datetime.now().isoformat()
    started_at = datetime.fromisoformat(st.session_state[state_key])
    deadline = started_at + timedelta(minutes=minutes_limit)
    seconds_left = max(0, int((deadline - datetime.now()).total_seconds()))
    return started_at, seconds_left


def format_seconds(seconds_left: int) -> str:
    """Format seconds into MM:SS."""
    minutes, seconds = divmod(max(0, seconds_left), 60)
    return f"{minutes:02d}:{seconds:02d}"


def get_student_identity() -> dict[str, str]:
    """Return the active authenticated student identity if present."""
    user = get_current_user()
    if user.get("role") == "student" and not user.get("is_guest"):
        return {
            "student_name": user.get("display_name", "").strip(),
            "student_key": user.get("email", "").strip().lower(),
        }
    return {"student_name": "", "student_key": ""}


def render_student_sign_in_panel(share_token: str) -> bool:
    """Render a compact sign-in panel for protected student links."""
    identity = get_student_identity()
    if identity["student_key"]:
        st.success(f"Signed in as {identity['student_name']} ({identity['student_key']}).")
        if st.button("Sign out student", key=f"student_share_signout_{share_token}", use_container_width=True):
            st.session_state.current_user = default_guest_user()
            st.rerun()
        return True

    st.warning("This test requires a student account. Sign in before starting the attempt.")
    with st.form(f"student_share_login_{share_token}", clear_on_submit=False):
        email = st.text_input("Student email")
        password = st.text_input("Student password", type="password")
        submitted = st.form_submit_button("Student Sign In", use_container_width=True)
    if submitted:
        user = authenticate_local_user(email, password)
        if user is None or user.get("role") != "student":
            st.error("This student account could not be verified. Use an existing student profile with the correct password.")
        else:
            user["is_guest"] = False
            st.session_state.current_user = user
            st.rerun()
    return False


def apply_student_draft_to_session(variant_data: dict[str, Any], variant_name: str, draft: dict[str, Any]) -> None:
    """Populate widget state from a saved student draft."""
    responses = draft.get("responses", {})
    for index, question in enumerate(variant_data.get("questions", [])):
        key_prefix = get_student_widget_prefix(variant_name, index)
        response = responses.get(f"question_{index}", "")
        if question["type"] in {"multiple_choice", "true_false", "short_answer"}:
            st.session_state[f"{key_prefix}_answer"] = response
        elif question["type"] == "matching" and isinstance(response, dict):
            for pair in question.get("pairs", []):
                st.session_state[f"{key_prefix}_match_{pair['left']}"] = response.get(pair["left"], "")


def maybe_autosave_student_draft(share_token: str, student_name: str, variant_data: dict[str, Any], variant_name: str) -> None:
    """Persist a draft only when the student has started answering and the content changed."""
    clean_name = student_name.strip()
    if not clean_name:
        return
    responses = collect_student_responses(variant_data, variant_name)
    if not any(str(value).strip() for value in responses.values() if not isinstance(value, dict)) and not any(
        any(str(item).strip() for item in value.values()) for value in responses.values() if isinstance(value, dict)
    ):
        return
    payload = {
        "variant_name": variant_name,
        "responses": responses,
    }
    signature = hashlib.sha256(json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest()
    session_key = f"draft_signature_{share_token}_{clean_name.lower()}"
    if st.session_state.get(session_key) == signature:
        return
    save_student_draft(share_token, clean_name, payload)
    st.session_state[session_key] = signature
    st.session_state[f"draft_saved_at_{share_token}_{clean_name.lower()}"] = datetime.now().strftime("%H:%M:%S")


def render_student_question(question: dict[str, Any], index: int, variant_name: str) -> None:
    """Render one question in student mode."""
    key_prefix = get_student_widget_prefix(variant_name, index)
    st.markdown(f"**Question {index + 1}. {question['question']}**")
    if question["type"] in {"multiple_choice", "true_false"}:
        st.radio(
            "Choose one answer",
            options=question.get("options", []),
            key=f"{key_prefix}_answer",
            label_visibility="collapsed",
        )
    elif question["type"] == "short_answer":
        st.text_input("Your answer", key=f"{key_prefix}_answer", label_visibility="collapsed")
    elif question["type"] == "matching":
        right_options = [pair["right"] for pair in question.get("pairs", [])]
        for pair in question.get("pairs", []):
            st.selectbox(
                f"{pair['left']}",
                options=[""] + right_options,
                key=f"{key_prefix}_match_{pair['left']}",
            )


def collect_student_responses(variant_data: dict[str, Any], variant_name: str) -> dict[str, Any]:
    """Collect student answers from Streamlit session state."""
    responses: dict[str, Any] = {}
    for index, question in enumerate(variant_data.get("questions", [])):
        key_prefix = get_student_widget_prefix(variant_name, index)
        if question["type"] in {"multiple_choice", "true_false", "short_answer"}:
            responses[f"question_{index}"] = st.session_state.get(f"{key_prefix}_answer", "")
        elif question["type"] == "matching":
            responses[f"question_{index}"] = {
                pair["left"]: st.session_state.get(f"{key_prefix}_match_{pair['left']}", "")
                for pair in question.get("pairs", [])
            }
    return responses


def save_attempt(
    variant_name: str,
    variant_data: dict[str, Any],
    student_name: str,
    result: dict[str, Any],
    share_token: str = "",
    submission_key: str = "",
    student_key: str = "",
) -> int:
    """Persist a student attempt and optionally sync it."""
    attempt_id = save_attempt_result(
        student_name=student_name,
        student_key=student_key,
        test_uid=get_current_test_uid(),
        variant_name=variant_name,
        test_title=variant_data.get("title", ""),
        owner_email=get_owner_email(),
        share_token=share_token,
        submission_key=submission_key,
        review_status="submitted",
        teacher_note="",
        answer_signature=str(result.get("attempt_meta", {}).get("answer_signature", "")),
        percentage=result["percentage"],
        payload=result,
    )
    maybe_sync_attempt(
        {
            "student_name": student_name,
            "test_uid": get_current_test_uid(),
            "variant_name": variant_name,
            "test_title": variant_data.get("title", ""),
            "percentage": result["percentage"],
            "owner_email": get_owner_email(),
            "payload": result,
        }
    )
    return attempt_id


def render_attempt_result(result: dict[str, Any]) -> None:
    """Render the latest student attempt analytics."""
    st.write("")
    st.subheader("Attempt Result")
    col1, col2, col3 = st.columns(3)
    col1.metric("Score", f"{result['total_score']}/{result['total_questions']}")
    col2.metric("Percentage", f"{result['percentage']}%")
    col3.metric("Questions", result["total_questions"])

    st.markdown("**Question-by-question feedback**")
    for item in result["per_question"]:
        if item["score"] >= 1.0:
            st.success(f"Q{item['index']}: Correct")
        else:
            st.error(f"Q{item['index']}: Not fully correct")
        st.caption(f"Skill: {item['skill_tag']}")
        st.write(f"Question: {item['question']}")
        st.write(f"Your answer: {item['student_answer']}")
        if item["correct_answer"]:
            st.write(f"Correct answer: {item['correct_answer']}")
        if item["explanation"]:
            st.write(f"Explanation: {item['explanation']}")


def render_student_submission_summary(result: dict[str, Any], student_name: str, show_score: bool) -> None:
    """Render a student-safe submission summary without revealing answers."""
    st.markdown(
        f"""
        <div class="hero-shell">
            <div class="hero-kicker">Submission received</div>
            <div class="hero-title" style="font-size: 1.8rem;">Thank you, {escape(student_name)}</div>
            <p class="hero-copy">Your answers were saved successfully. The teacher can now review your attempt.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if not show_score:
        st.info("Your response has been submitted successfully.")
        return

    col1, col2, col3 = st.columns(3)
    col1.metric("Score", f"{result['total_score']}/{result['total_questions']}")
    col2.metric("Percentage", f"{result['percentage']}%")
    col3.metric("Answered", result["total_questions"])
    st.caption("Correct answers and explanations are hidden in student mode.")


def get_share_token_from_query() -> str:
    """Read a share token from query params if present."""
    value = st.query_params.get("share", "")
    if isinstance(value, list):
        return value[0] if value else ""
    return str(value).strip()


def save_shared_attempt(
    *,
    shared_record: dict[str, Any],
    student_name: str,
    student_key: str,
    result: dict[str, Any],
) -> int:
    """Save an attempt submitted through a public share link."""
    variant_data = shared_record["payload"]["variant_data"]
    responses = result.get("responses", {})
    submission_key = build_submission_key(shared_record["token"], student_name, responses)
    attempt_id = save_attempt_result(
        student_name=student_name,
        student_key=student_key,
        test_uid=shared_record.get("test_uid", ""),
        variant_name=shared_record["variant_name"],
        test_title=shared_record["title"],
        owner_email=shared_record.get("owner_email", ""),
        share_token=shared_record["token"],
        submission_key=submission_key,
        review_status="submitted",
        teacher_note="",
        answer_signature=str(result.get("attempt_meta", {}).get("answer_signature", "")),
        percentage=result["percentage"],
        payload=result,
    )
    maybe_sync_attempt(
        {
            "student_name": student_name,
            "test_uid": shared_record.get("test_uid", ""),
            "variant_name": shared_record["variant_name"],
            "test_title": shared_record["title"],
            "percentage": result["percentage"],
            "owner_email": shared_record.get("owner_email", ""),
            "payload": result,
        }
    )
    return attempt_id


def render_submission_success_card(student_name: str, result: dict[str, Any]) -> None:
    """Render a cleaner success screen after student submission."""
    render_student_submission_summary(result, student_name, show_score=True)


def render_shared_student_page(share_token: str) -> None:
    """Render the public student page for a shared test link."""
    shared_record = load_share_link(share_token)
    if shared_record is None:
        st.error("This shared test link does not exist.")
        return
    if not shared_record.get("is_active"):
        st.warning("This shared test link is inactive.")
        return

    deadline = parse_deadline(shared_record.get("deadline_at", ""))
    if deadline and datetime.now() > deadline:
        st.warning("This test is closed because the deadline has passed.")
        return

    variant_data = shared_record["payload"]["variant_data"]
    variant_name = shared_record["variant_name"]
    share_settings = shared_record.get("payload", {}).get("share_settings", {})
    require_student_login = bool(share_settings.get("require_student_login", False))
    reveal_score_after_submit = bool(share_settings.get("reveal_score_after_submit", True))
    allowed_students = parse_whitelist("\n".join(share_settings.get("allowed_students", [])))
    per_student_random_order = bool(share_settings.get("per_student_random_order", False))
    timer_minutes = int(share_settings.get("timer_minutes", 0) or 0)
    one_question_at_a_time = bool(share_settings.get("one_question_at_a_time", False))
    block_copy_print = bool(share_settings.get("block_copy_print", False))
    st.markdown(
        f"""
        <div class="hero-shell student-shell">
            <div class="hero-kicker">Student Assessment</div>
            <div class="hero-title" style="font-size: 1.9rem;">{escape(shared_record['title'])}</div>
            <p class="hero-copy">Variant {escape(variant_name)}. Your progress is saved automatically. Review each answer carefully before the final submission.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if variant_data.get("instructions"):
        st.info(variant_data["instructions"])
    if deadline:
        st.caption(f"Deadline: {deadline.strftime('%Y-%m-%d %H:%M')}")

    success_key = f"shared_success_{share_token}"
    if success_key in st.session_state:
        render_student_submission_summary(
            st.session_state[success_key]["result"],
            st.session_state[success_key]["student_name"],
            reveal_score_after_submit,
        )
        return

    if require_student_login and not render_student_sign_in_panel(share_token):
        return

    identity = get_student_identity()
    student_name_key = f"shared_student_name_{share_token}"
    if require_student_login:
        student_name = identity["student_name"]
        student_key = identity["student_key"]
        st.text_input(
            "Student name",
            value=student_name,
            key=student_name_key,
            disabled=True,
        )
    else:
        student_name = st.text_input("Student name", key=student_name_key, placeholder="Enter your full name")
        student_key = student_name.strip().lower()

    if require_student_login and allowed_students and student_key not in allowed_students:
        st.error("This student account is not on the allowed list for this test.")
        return
    if require_student_login and count_share_attempts_for_student_key(shared_record["token"], student_key) >= 1:
        st.warning("This student account has already submitted this test. Re-entry is blocked.")
        return

    personalized_variant = build_personalized_variant(
        variant_data,
        share_token=share_token,
        student_identity=student_key or student_name.strip(),
        enable_random_order=per_student_random_order,
    )
    watermark_text = student_name or student_key or "Student attempt"
    render_soft_exam_protection(watermark_text, block_copy_print)

    draft_identity = student_key or student_name.strip()
    load_key = f"loaded_draft_{share_token}_{draft_identity.lower()}"
    if student_name.strip() and not st.session_state.get(load_key):
        draft = load_student_draft(share_token, draft_identity)
        if draft is not None:
            apply_student_draft_to_session(personalized_variant, variant_name, draft)
        st.session_state[load_key] = True

    total_questions = max(1, len(personalized_variant.get("questions", [])))
    answered_questions = count_completed_answers(personalized_variant, variant_name)
    progress = answered_questions / total_questions

    _, seconds_left = get_exam_timer_state(share_token, draft_identity, timer_minutes)
    if timer_minutes > 0 and seconds_left <= 0:
        st.error("Time is over. This test has been closed automatically.")
        return

    metric_col1, metric_col2, metric_col3, metric_col4 = st.columns(4, gap="large")
    metric_col1.metric("Completed", f"{answered_questions}/{total_questions}")
    metric_col2.metric("Variant", variant_name)
    metric_col3.metric("Status", "Ready to submit" if answered_questions == total_questions else "In progress")
    metric_col4.metric("Time left", format_seconds(seconds_left) if timer_minutes > 0 else "No timer")
    helper_col1, helper_col2 = st.columns(2, gap="large")
    helper_col1.caption("Student page is optimized for phones: focused layout, fewer distractions, and automatic draft saving.")
    helper_col2.caption("Correct answers stay hidden in student mode, even after submission.")
    if timer_minutes > 0:
        st.caption("The timer is enforced on submit and while navigating the test, without forcing a full page reload.")
    st.progress(progress, text=f"Progress: {answered_questions} of {total_questions} answered")

    if one_question_at_a_time:
        current_index_key = f"student_page_index_{share_token}_{draft_identity.lower()}"
        current_index = int(st.session_state.get(current_index_key, 0))
        current_index = max(0, min(current_index, total_questions - 1))
        question = personalized_variant.get("questions", [])[current_index]
        st.caption(f"Question {current_index + 1} of {total_questions}")
        with st.container(border=True):
            render_student_question(question, current_index, variant_name)
        nav_col1, nav_col2 = st.columns(2, gap="large")
        with nav_col1:
            if st.button("Previous", disabled=current_index == 0, use_container_width=True, key=f"prev_q_{share_token}"):
                st.session_state[current_index_key] = max(0, current_index - 1)
                st.rerun()
        with nav_col2:
            next_label = "Next" if current_index < total_questions - 1 else "Review answers"
            if st.button(next_label, use_container_width=True, key=f"next_q_{share_token}"):
                st.session_state[current_index_key] = min(total_questions - 1, current_index + 1)
                st.rerun()
    else:
        for index, question in enumerate(personalized_variant.get("questions", [])):
            with st.container(border=True):
                render_student_question(question, index, variant_name)
            st.write("")

    maybe_autosave_student_draft(share_token, draft_identity, personalized_variant, variant_name)

    action_col1, action_col2 = st.columns([1, 1], gap="large")
    with action_col1:
        if student_name.strip():
            st.caption("Your draft is saved automatically on this device and in the local app database.")
            saved_at = st.session_state.get(f"draft_saved_at_{share_token}_{draft_identity.lower()}", "")
            if saved_at:
                st.caption(f"Saved just now: {saved_at}")
        else:
            st.caption("Enter your name first to enable draft saving.")
    with action_col2:
        if one_question_at_a_time:
            st.caption("Questions are shown one at a time to reduce copying and answer sharing.")
        else:
            st.caption("Use the confirmation block below before the final submission.")
        st.caption("If your connection drops, reopen the same link. Your draft and session state will be restored when available.")

    st.markdown("**Finish test**")
    confirm_ready = st.checkbox(
        "I have reviewed my answers and I am ready to submit this test.",
        key=f"confirm_ready_{share_token}",
    )
    submit_clicked = st.button("Finish Test", use_container_width=True, type="primary")

    if submit_clicked:
        if not student_name.strip():
            st.error("Please enter your name before submitting.")
            return
        if not confirm_ready:
            st.warning("Please confirm that you are ready to submit.")
            return
        max_attempts = int(shared_record.get("max_attempts", 1))
        if max_attempts != 0:
            if require_student_login:
                current_attempts = count_share_attempts_for_student_key(shared_record["token"], student_key)
                if current_attempts >= 1:
                    st.error("This student account has already submitted this test. A second attempt is not allowed.")
                    return
            else:
                current_attempts = count_share_attempts(shared_record["token"], student_name.strip())
                if current_attempts >= max_attempts:
                    st.error("This student has already reached the allowed attempt limit.")
                    return
        responses = collect_student_responses(personalized_variant, variant_name)
        submission_key = build_submission_key(shared_record["token"], student_name.strip(), responses)
        if attempt_submission_exists(submission_key):
            st.warning("These exact answers were already submitted. The teacher page already has this attempt.")
            return
        result = grade_attempt(personalized_variant, responses)
        result["responses"] = responses
        started_at, _ = get_exam_timer_state(share_token, draft_identity, timer_minutes)
        duration_seconds = max(0, int((datetime.now() - started_at).total_seconds()))
        result["attempt_meta"] = {
            "duration_seconds": duration_seconds,
            "submitted_at": datetime.now().isoformat(),
            "timer_minutes": timer_minutes,
            "one_question_at_a_time": one_question_at_a_time,
            "answer_signature": build_answer_signature(responses),
        }
        save_shared_attempt(
            shared_record=shared_record,
            student_name=student_name.strip(),
            student_key=student_key,
            result=result,
        )
        delete_student_draft(share_token, draft_identity)
        st.session_state[success_key] = {
            "student_name": student_name.strip(),
            "result": result,
        }
        st.rerun()


def render_student_mode(variants: dict[str, dict[str, Any]], disable: bool) -> None:
    """Render student practice mode with automatic checking."""
    st.subheader("Student Practice Mode")
    if disable:
        st.info("Student mode is locked until the test passes the quality checks.")
        return

    variant_name = st.selectbox("Choose variant", options=list(variants), key="student_variant")
    student_name = st.text_input("Student name", value="Student 1", key="student_name")
    variant_data = variants[variant_name]

    answered_questions = count_completed_answers(variant_data, variant_name)
    total_questions = max(1, len(variant_data.get("questions", [])))
    st.progress(answered_questions / total_questions, text=f"Progress: {answered_questions}/{total_questions}")

    for index, question in enumerate(variant_data.get("questions", [])):
        with st.container(border=True):
            render_student_question(question, index, variant_name)
        st.write("")

    confirm_submit = st.checkbox(
        "I have reviewed my answers before submitting.",
        key=f"practice_confirm_{variant_name}",
    )
    submitted = st.button("Submit Answers", use_container_width=True, key=f"practice_submit_{variant_name}")

    if submitted:
        if not confirm_submit:
            st.warning("Please confirm the submission first.")
            return
        responses = collect_student_responses(variant_data, variant_name)
        result = grade_attempt(variant_data, responses)
        result["responses"] = responses
        result["attempt_meta"] = {
            "duration_seconds": 0,
            "submitted_at": datetime.now().isoformat(),
            "timer_minutes": 0,
            "one_question_at_a_time": False,
            "answer_signature": build_answer_signature(responses),
        }
        attempt_id = save_attempt(variant_name, variant_data, student_name, result)
        st.session_state.last_attempt = {
            "attempt_id": attempt_id,
            "student_name": student_name,
            "result": result,
        }

    if st.session_state.last_attempt:
        render_attempt_result(st.session_state.last_attempt["result"])


def build_attempt_export_frames(attempts: list[dict[str, Any]]) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Build summary and per-question frames for analytics export."""
    summary_df = pd.DataFrame(
        [
            {
                "Attempt ID": item["id"],
                "Student": item["student_name"],
                "Variant": item["variant_name"],
                "Test": item["test_title"],
                "Score %": item["percentage"],
                "Status": item.get("review_status", "submitted"),
                "Teacher Note": item.get("teacher_note", ""),
                "Share Token": item.get("share_token", ""),
                "Submitted At": item["created_at"],
            }
            for item in attempts
        ]
    )
    detail_rows: list[dict[str, Any]] = []
    for item in attempts:
        for detail in item.get("details", {}).get("per_question", []):
            detail_rows.append(
                {
                    "Attempt ID": item["id"],
                    "Student": item["student_name"],
                    "Variant": item["variant_name"],
                    "Question #": detail.get("index", 0),
                    "Question": detail.get("question", ""),
                    "Skill": detail.get("skill_tag", ""),
                    "Student Answer": detail.get("student_answer", ""),
                    "Correct Answer": detail.get("correct_answer", ""),
                    "Score": detail.get("score", 0),
                }
            )
    return summary_df, pd.DataFrame(detail_rows)


def render_analytics_export(attempts: list[dict[str, Any]]) -> None:
    """Render CSV/XLSX analytics exports."""
    summary_df, details_df = build_attempt_export_frames(attempts)
    csv_bytes = summary_df.to_csv(index=False).encode("utf-8")
    xlsx_buffer = BytesIO()
    with pd.ExcelWriter(xlsx_buffer, engine="openpyxl") as writer:
        summary_df.to_excel(writer, sheet_name="Summary", index=False)
        if not details_df.empty:
            details_df.to_excel(writer, sheet_name="Details", index=False)
    xlsx_buffer.seek(0)

    export_col1, export_col2 = st.columns(2, gap="large")
    with export_col1:
        st.download_button(
            "Export analytics CSV",
            csv_bytes,
            "analytics_summary.csv",
            "text/csv",
            use_container_width=True,
        )
    with export_col2:
        st.download_button(
            "Export analytics Excel",
            xlsx_buffer.getvalue(),
            "analytics_summary.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )


def build_backup_bundle() -> bytes:
    """Build a JSON backup for the current teacher account."""
    bundle = {
        "exported_at": datetime.now().isoformat(),
        "owner_email": get_owner_email(),
        "tests": list_test_history(limit=1000, owner_email=get_owner_email()),
        "question_bank": list_question_bank(limit=1000, owner_email=get_owner_email()),
        "attempts": list_attempt_results(limit=5000, owner_email=get_owner_email()),
        "share_links": list_share_links(limit=1000, owner_email=get_owner_email()),
        "groups": list_student_groups(get_owner_email()),
        "roster": list_group_students(get_owner_email()),
        "api_errors": list_api_error_logs(limit=200),
    }
    return json.dumps(bundle, ensure_ascii=False, indent=2).encode("utf-8")


def normalize_student_import_frame(frame: pd.DataFrame) -> list[dict[str, Any]]:
    """Normalize imported student spreadsheets into a stable schema."""
    renamed = {str(column).strip().lower(): str(column) for column in frame.columns}
    normalized_rows: list[dict[str, Any]] = []
    for _, row in frame.iterrows():
        values = {str(key).strip().lower(): row[key] for key in frame.columns}
        normalized_rows.append(
            {
                "full_name": str(values.get("full_name", values.get("name", values.get("student", ""))) or "").strip(),
                "email": str(values.get("email", "") or "").strip().lower(),
                "external_id": str(values.get("external_id", values.get("student_id", values.get("id", ""))) or "").strip(),
                "notes": str(values.get("notes", "") or "").strip(),
            }
        )
    return normalized_rows


def render_backup_center() -> None:
    """Render teacher backup utilities."""
    open_section("Backup Center")
    st.caption("Export a full teacher backup before live demos, major refactors, or data migrations.")
    backup_bytes = build_backup_bundle()
    st.download_button(
        "Download full backup JSON",
        backup_bytes,
        f"teacher_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
        "application/json",
        use_container_width=True,
    )
    close_section()


def render_groups_and_roster_view() -> None:
    """Render group management and student import tools."""
    open_section("Groups, Classes, and Student Import")
    groups = list_student_groups(get_owner_email())
    group_options = {f"{item['name']} ({item.get('grade_level', 'No grade')})": int(item["id"]) for item in groups}

    creation_col, import_col = st.columns([1.1, 1.4], gap="large")
    with creation_col:
        st.markdown("**Create class or group**")
        with st.form("create_group_form", clear_on_submit=True):
            group_name = st.text_input("Group name", placeholder="5A Mathematics")
            group_grade = st.text_input("Grade / class", placeholder="5-6 grade")
            group_description = st.text_area("Description", placeholder="Morning class, algebra focus", height=90)
            create_group_submitted = st.form_submit_button("Create group", use_container_width=True)
        if create_group_submitted:
            if not group_name.strip():
                st.error("Enter a group name before creating the class.")
            elif len(group_name.strip()) < 2:
                st.error("Group name should contain at least 2 characters.")
            else:
                group_id = create_student_group(
                    owner_email=get_owner_email(),
                    name=group_name,
                    grade_level=group_grade,
                    description=group_description,
                )
                log_event("create_group", "group", str(group_id), {"name": group_name.strip(), "grade_level": group_grade.strip()})
                st.success("Group created successfully.")
                st.rerun()

        if groups:
            st.markdown("**Add one student manually**")
            selected_manual_group = st.selectbox("Target group", options=list(group_options), key="manual_group_pick")
            with st.form("manual_student_form", clear_on_submit=True):
                student_name = st.text_input("Student full name")
                student_email = st.text_input("Student email")
                student_external_id = st.text_input("Student ID (optional)")
                student_notes = st.text_input("Notes (optional)")
                save_student_submitted = st.form_submit_button("Add student", use_container_width=True)
            if save_student_submitted:
                if not student_name.strip():
                    st.error("Student full name is required.")
                elif student_email.strip() and not is_valid_email(student_email):
                    st.error("Student email format is invalid.")
                else:
                    allowed, limit_message = can_add_students(1)
                    if not allowed:
                        st.error(limit_message)
                    else:
                        save_group_student(
                            owner_email=get_owner_email(),
                            group_id=group_options[selected_manual_group],
                            full_name=student_name,
                            email=student_email,
                            external_id=student_external_id,
                            notes=student_notes,
                        )
                        record_usage_event(get_owner_email(), "student_import", 1, {"mode": "manual"})
                        log_event("add_student", "group", str(group_options[selected_manual_group]), {"student_email": student_email.strip().lower(), "student_name": student_name.strip()})
                        st.success("Student added to the group.")
                        st.rerun()

    with import_col:
        st.markdown("**Import student list**")
        if not groups:
            st.info("Create at least one group before importing a student list.")
        else:
            selected_import_group = st.selectbox("Import into group", options=list(group_options), key="import_group_pick")
            roster_file = st.file_uploader(
                "Upload CSV or Excel roster",
                type=["csv", "xlsx"],
                key="roster_import_file",
                help="Columns can be named full_name/name, email, external_id/student_id, notes.",
            )
            pasted_rows = st.text_area(
                "Or paste roster lines",
                placeholder="Aruzhan Nur\nDias Bektur, dias@example.com",
                height=100,
                key="roster_paste_input",
            )
            if st.button("Import students", use_container_width=True, key="roster_import_button"):
                rows: list[dict[str, Any]] = []
                if roster_file is not None:
                    frame = pd.read_csv(roster_file) if roster_file.name.lower().endswith(".csv") else pd.read_excel(roster_file)
                    rows.extend(normalize_student_import_frame(frame))
                if pasted_rows.strip():
                    for line in pasted_rows.splitlines():
                        parts = [part.strip() for part in line.split(",")]
                        if not parts or not parts[0]:
                            continue
                        rows.append(
                            {
                                "full_name": parts[0],
                                "email": parts[1].lower() if len(parts) > 1 else "",
                                "external_id": parts[2] if len(parts) > 2 else "",
                                "notes": "",
                            }
                        )
                if not rows:
                    st.error("Upload a roster file or paste student rows before importing.")
                else:
                    invalid_emails = [row["email"] for row in rows if row.get("email") and not is_valid_email(row["email"])]
                    if invalid_emails:
                        st.error("Some student emails are invalid. Fix them before importing.")
                    else:
                        allowed, limit_message = can_add_students(len(rows))
                        if not allowed:
                            st.error(limit_message)
                        else:
                            imported = import_group_students(
                                owner_email=get_owner_email(),
                                group_id=group_options[selected_import_group],
                                rows=rows,
                            )
                            record_usage_event(get_owner_email(), "student_import", imported, {"mode": "bulk"})
                            log_event("import_students", "group", str(group_options[selected_import_group]), {"count": imported})
                            st.success(f"Imported or updated {imported} student records.")
                            st.rerun()

    roster = list_group_students(get_owner_email())
    if roster:
        st.markdown("**Current roster**")
        st.dataframe(pd.DataFrame(roster), use_container_width=True, hide_index=True)
    else:
        st.info("No students imported yet.")
    close_section()


def render_gradebook_view() -> None:
    """Render an electronic gradebook with teacher-friendly summaries."""
    attempts = list_attempt_results(limit=2000, owner_email=get_owner_email())
    roster = get_owner_roster()
    open_section("Electronic Gradebook")
    if not attempts:
        st.info("The gradebook will appear after student submissions.")
        close_section()
        return

    gradebook_rows = build_gradebook_rows(attempts, roster)
    gradebook_df = pd.DataFrame(gradebook_rows)
    filter_col1, filter_col2, filter_col3 = st.columns(3, gap="large")
    with filter_col1:
        group_filter = st.selectbox(
            "Filter by group",
            options=["All"] + sorted({row.get("Group", "") for row in gradebook_rows if row.get("Group", "")}),
            key="gradebook_group_filter",
        )
    with filter_col2:
        risk_filter = st.selectbox("Risk band", options=["All", "Critical", "High", "Moderate", "Low"], key="gradebook_risk_filter")
    with filter_col3:
        search_filter = st.text_input("Search student", key="gradebook_search")

    filtered_df = gradebook_df.copy()
    if group_filter != "All":
        filtered_df = filtered_df[filtered_df["Group"] == group_filter]
    if risk_filter != "All":
        filtered_df = filtered_df[filtered_df["Risk"] == risk_filter]
    if search_filter.strip():
        filtered_df = filtered_df[filtered_df["Student"].str.contains(search_filter.strip(), case=False, na=False)]

    st.dataframe(filtered_df, use_container_width=True, hide_index=True)

    gradebook_buffer = BytesIO()
    with pd.ExcelWriter(gradebook_buffer, engine="openpyxl") as writer:
        filtered_df.to_excel(writer, sheet_name="Gradebook", index=False)
    gradebook_buffer.seek(0)
    export_col1, export_col2 = st.columns(2, gap="large")
    with export_col1:
        st.download_button(
            "Export gradebook CSV",
            filtered_df.to_csv(index=False).encode("utf-8"),
            "gradebook.csv",
            "text/csv",
            use_container_width=True,
        )
    with export_col2:
        st.download_button(
            "Export gradebook Excel",
            gradebook_buffer.getvalue(),
            "gradebook.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
    close_section()


def render_teacher_home_dashboard() -> None:
    """Render one home dashboard for daily teacher operations."""
    open_section("Teacher Command Center")
    plan = get_current_plan_status()
    today_prefix = datetime.now().strftime("%Y-%m-%d")
    current_test_uid = get_current_test_uid()
    attempts_current = list_attempt_results(limit=500, owner_email=get_owner_email(), test_uid=current_test_uid or None)
    all_attempts = list_attempt_results(limit=1000, owner_email=get_owner_email())
    active_links = [item for item in list_share_links(limit=200, owner_email=get_owner_email()) if bool(item.get("is_active"))]
    todays_attempts = [item for item in all_attempts if str(item.get("created_at", "")).startswith(today_prefix)]
    suspicious = detect_suspicious_attempts(attempts_current)
    weak_topics = build_topic_progress_rows(attempts_current).get("overall", [])[:5]
    pending_reviews = [item for item in attempts_current if str(item.get("review_status", "submitted")) == "submitted"]

    top1, top2, top3, top4, top5 = st.columns(5, gap="large")
    top1.metric("Plan", str(plan.get("plan_name", "free")).replace("_", " ").title())
    top2.metric("This month", f"{plan['usage'].get('monthly_generations', 0)}/{plan['limits'].get('monthly_generations', 0)}")
    top3.metric("Today", len(todays_attempts))
    top4.metric("Active links", len(active_links))
    top5.metric("Pending reviews", len(pending_reviews))

    left, right = st.columns([1.2, 1], gap="large")
    with left:
        st.markdown("**Today and next actions**")
        if pending_reviews:
            st.warning(f"{len(pending_reviews)} attempts still need review.")
        else:
            st.success("No pending reviews right now.")
        if weak_topics:
            st.markdown("**Weak topics right now**")
            st.dataframe(pd.DataFrame(weak_topics), use_container_width=True, hide_index=True)
        else:
            st.info("Weak-topic signals will appear after more submissions.")
    with right:
        st.markdown("**Suspicious attempts**")
        suspicious_df = pd.DataFrame(suspicious)
        if not suspicious_df.empty:
            st.dataframe(suspicious_df.head(5), use_container_width=True, hide_index=True)
        else:
            st.success("No suspicious attempts flagged for the current test.")
    close_section()


def render_business_center() -> None:
    """Render pricing, usage, migration, and audit tools."""
    open_section("Business, Billing, and Operations")
    plan = get_current_plan_status()
    pricing_rows = [
        {"Plan": "Free", "Monthly generations": 30, "Students": 50, "Active tests": 10, "Best for": "Small experiments"},
        {"Plan": "Teacher Pro", "Monthly generations": 500, "Students": 1000, "Active tests": 200, "Best for": "Independent teachers"},
        {"Plan": "School", "Monthly generations": 5000, "Students": 10000, "Active tests": 5000, "Best for": "Departments and schools"},
    ]
    usage_rows = list_usage_events(limit=200, owner_email=get_owner_email())
    audit_rows = list_audit_logs(limit=100, actor_email=get_owner_email())

    metric_col1, metric_col2, metric_col3 = st.columns(3, gap="large")
    metric_col1.metric("Current plan", str(plan.get("plan_name", "free")).replace("_", " ").title())
    metric_col2.metric("Generation usage", f"{plan['usage'].get('monthly_generations', 0)} / {plan['limits'].get('monthly_generations', 0)}")
    metric_col3.metric("Trial ends", plan.get("trial_ends_at", "")[:10] or "N/A")

    info_col1, info_col2 = st.columns(2, gap="large")
    with info_col1:
        st.markdown("**Pricing table**")
        st.dataframe(pd.DataFrame(pricing_rows), use_container_width=True, hide_index=True)
        st.caption("Stripe is not wired yet, but plans, quotas, and dashboards are now present in the product layer.")
    with info_col2:
        st.markdown("**Migration and recovery**")
        if is_cloud_enabled():
            if st.button("Migrate local teacher data to Supabase", use_container_width=True):
                migrated = migrate_local_data_to_cloud(get_owner_email())
                log_event("migrate_local_to_cloud", "migration", get_owner_email(), migrated)
                st.success(
                    "Migration finished: "
                    f"{migrated['users']} users, {migrated['tests']} tests, {migrated['attempts']} attempts, "
                    f"{migrated['groups']} groups, {migrated['students']} students."
                )
        else:
            st.info("Enable Supabase first to migrate local data and use cloud recovery.")
        st.caption("Recovery flow: the app already falls back to SQLite/local generator when cloud services are unavailable.")

    lower1, lower2 = st.columns(2, gap="large")
    with lower1:
        st.markdown("**Recent usage events**")
        if usage_rows:
            st.dataframe(pd.DataFrame(usage_rows), use_container_width=True, hide_index=True)
        else:
            st.info("Usage events will appear after generation, imports, and other actions.")
    with lower2:
        st.markdown("**Recent audit log**")
        if audit_rows:
            st.dataframe(pd.DataFrame(audit_rows), use_container_width=True, hide_index=True)
        else:
            st.info("Audit events will appear after sign-in, sharing, edits, and admin actions.")
    close_section()


def build_student_weak_topics(student_attempts: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Build a compact weak-topic summary for one student."""
    topic_errors: dict[str, int] = defaultdict(int)
    topic_totals: dict[str, int] = defaultdict(int)
    for attempt in student_attempts:
        for item in attempt.get("details", {}).get("per_question", []):
            skill = item.get("skill_tag", "") or "General"
            topic_totals[skill] += 1
            if float(item.get("score", 0.0)) < 0.999:
                topic_errors[skill] += 1
    rows = []
    for skill, total in topic_totals.items():
        mistakes = topic_errors.get(skill, 0)
        accuracy = round(((total - mistakes) / total) * 100, 2) if total else 0.0
        rows.append(
            {
                "Skill": skill,
                "Mistakes": mistakes,
                "Attempts": total,
                "Accuracy %": accuracy,
                "Risk": classify_risk(accuracy),
            }
        )
    rows.sort(key=lambda row: (row["Accuracy %"], -row["Mistakes"]))
    return rows[:5]


def render_analytics_dashboard() -> None:
    """Render detailed analytics for the currently opened test."""
    st.subheader("Analytics Dashboard")
    attempts = list_attempt_results(limit=200, owner_email=get_owner_email(), test_uid=get_current_test_uid())
    all_attempts = list_attempt_results(limit=1000, owner_email=get_owner_email())
    aggregate = aggregate_attempt_history(attempts)
    topic_progress = build_topic_progress_rows(attempts)
    suspicious_rows = detect_suspicious_attempts(attempts)

    st.caption("This dashboard is scoped to the current test only.")

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Saved attempts", aggregate["attempt_count"])
    col2.metric("Average score", f"{aggregate['average_percentage']}%")
    col3.metric("Median score", f"{aggregate['median_percentage']}%")
    col4.metric("Pass rate", f"{aggregate['pass_rate']}%")

    if aggregate["attempt_count"] == 0:
        st.info("No student attempts yet for this test.")
        return

    render_analytics_export(attempts)

    info_col1, info_col2 = st.columns(2)
    info_col1.metric("Unique students", aggregate["unique_students"])
    variant_gap = 0.0
    if aggregate["variant_performance"]:
        variant_gap = round(max(aggregate["variant_performance"].values()) - min(aggregate["variant_performance"].values()), 2)
    info_col2.metric("Variant gap", f"{variant_gap}%")

    risk_col, recommendation_col = st.columns(2, gap="large")
    with risk_col:
        st.markdown("**Risk Alerts**")
        if aggregate["risk_alerts"]:
            for item in aggregate["risk_alerts"]:
                st.warning(item)
        else:
            st.success("No major risks detected for this test.")
    with recommendation_col:
        st.markdown("**Recommended Actions**")
        if aggregate["recommendations"]:
            for item in aggregate["recommendations"]:
                st.info(item)
        else:
            st.info("No intervention recommendations yet.")

    if aggregate["by_type"]:
        st.markdown("**Accuracy by question type**")
        type_df = pd.DataFrame(
            {
                "Question Type": list(aggregate["by_type"].keys()),
                "Accuracy": list(aggregate["by_type"].values()),
            }
        ).set_index("Question Type")
        st.bar_chart(type_df)

    if aggregate["variant_performance"]:
        st.markdown("**Performance by variant**")
        variant_df = pd.DataFrame(
            {
                "Variant": list(aggregate["variant_performance"].keys()),
                "Average Score %": list(aggregate["variant_performance"].values()),
            }
        ).set_index("Variant")
        st.bar_chart(variant_df)

    if aggregate["error_topics"]:
        st.markdown("**Topics or skills with the most mistakes**")
        topic_df = pd.DataFrame(
            {
                "Skill or Topic": list(aggregate["error_topics"].keys()),
                "Mistakes": list(aggregate["error_topics"].values()),
            }
        )
        st.dataframe(topic_df, use_container_width=True, hide_index=True)

    progress_col1, progress_col2 = st.columns(2, gap="large")
    with progress_col1:
        st.markdown("**Topic progress overview**")
        overall_topic_df = pd.DataFrame(topic_progress["overall"])
        if not overall_topic_df.empty:
            st.dataframe(overall_topic_df, use_container_width=True, hide_index=True)
        else:
            st.info("Topic mastery will appear after more question-level data is collected.")
    with progress_col2:
        st.markdown("**Suspicious attempts**")
        suspicious_df = pd.DataFrame(suspicious_rows)
        if not suspicious_df.empty:
            st.dataframe(suspicious_df, use_container_width=True, hide_index=True)
        else:
            st.success("No suspicious patterns were flagged for this test.")

    if aggregate["weak_topics_priority"]:
        st.markdown("**Weak topics by priority**")
        st.dataframe(pd.DataFrame(aggregate["weak_topics_priority"]), use_container_width=True, hide_index=True)

    st.markdown("**Student cards**")
    student_profiles = aggregate["student_profiles"]
    if student_profiles:
        for student_profile in student_profiles:
            student_name = student_profile["Student"]
            student_attempts_current = [item for item in attempts if item["student_name"] == student_name]
            student_attempts_all = [item for item in all_attempts if item["student_name"] == student_name]
            gap_vs_average = round(float(student_profile["Average %"]) - float(aggregate["average_percentage"]), 2)
            weak_topics = build_student_weak_topics(student_attempts_current)
            with st.container(border=True):
                title_col, metric_col = st.columns([2.8, 1.2], gap="large")
                with title_col:
                    st.markdown(f"**{student_name}**")
                    st.caption("Personal analytics card for the current test, with progress across all saved tests.")
                with metric_col:
                    st.metric("Vs class average", f"{gap_vs_average:+.2f}%")

                meta_col1, meta_col2, meta_col3, meta_col4 = st.columns(4, gap="large")
                meta_col1.metric("Current avg", f"{student_profile['Average %']}%")
                meta_col2.metric("Best", f"{student_profile['Best %']}%")
                meta_col3.metric("Attempts", student_profile["Attempts"])
                meta_col4.metric("Risk", student_profile["Risk"])

                detail_col1, detail_col2 = st.columns([1.2, 1], gap="large")
                with detail_col1:
                    st.markdown("**Weak topics for this student**")
                    if weak_topics:
                        st.dataframe(pd.DataFrame(weak_topics), use_container_width=True, hide_index=True)
                    else:
                        st.info("No weak-topic signals yet.")
                with detail_col2:
                    st.markdown("**Progress across multiple tests**")
                    progress_rows = [
                        {
                            "Date": item["created_at"][:10],
                            "Test": item["test_title"],
                            "Score %": item["percentage"],
                        }
                        for item in reversed(student_attempts_all)
                    ]
                    if progress_rows:
                        progress_df = pd.DataFrame(progress_rows)
                        st.dataframe(progress_df, use_container_width=True, hide_index=True)
                        timeline_df = progress_df[["Date", "Score %"]].copy().set_index("Date")
                        st.line_chart(timeline_df)
                    else:
                        st.info("No multi-test history yet.")

    insight_col1, insight_col2 = st.columns(2, gap="large")
    with insight_col1:
        st.markdown("**Skill Risk Table**")
        if aggregate["skill_insights"]:
            st.dataframe(pd.DataFrame(aggregate["skill_insights"]), use_container_width=True, hide_index=True)
        else:
            st.info("No skill data yet.")
    with insight_col2:
        st.markdown("**Student Risk Table**")
        if aggregate["student_risks"]:
            st.dataframe(pd.DataFrame(aggregate["student_risks"]), use_container_width=True, hide_index=True)
        else:
            st.info("No student risk data yet.")

    student_detail_options = [row["Student"] for row in aggregate["student_profiles"]]
    if student_detail_options:
        selected_student = st.selectbox("Open individual student analytics", options=student_detail_options)
        student_attempts = [item for item in attempts if item["student_name"] == selected_student]
        student_attempts_all = [item for item in all_attempts if item["student_name"] == selected_student]
        student_profile = next((row for row in aggregate["student_profiles"] if row["Student"] == selected_student), None)
        if student_profile:
            student_col1, student_col2, student_col3, student_col4 = st.columns(4)
            student_col1.metric("Average %", student_profile["Average %"])
            student_col2.metric("Attempts", student_profile["Attempts"])
            student_col3.metric("Risk", student_profile["Risk"])
            student_col4.metric("Vs class average", f"{round(float(student_profile['Average %']) - float(aggregate['average_percentage']), 2):+.2f}%")
            student_attempt_df = pd.DataFrame(
                [
                    {
                        "Variant": item["variant_name"],
                        "Score %": item["percentage"],
                        "Submitted At": item["created_at"],
                    }
                    for item in student_attempts
                ]
            )
            st.dataframe(student_attempt_df, use_container_width=True, hide_index=True)
            weak_topics_df = pd.DataFrame(build_student_weak_topics(student_attempts))
            if not weak_topics_df.empty:
                st.markdown("**Weak topics for selected student**")
                st.dataframe(weak_topics_df, use_container_width=True, hide_index=True)
            if student_attempts_all:
                all_tests_df = pd.DataFrame(
                    [
                        {
                            "Date": item["created_at"][:10],
                            "Test": item["test_title"],
                            "Variant": item["variant_name"],
                            "Score %": item["percentage"],
                        }
                        for item in reversed(student_attempts_all)
                    ]
                )
                st.markdown("**Progress across all saved tests**")
                st.dataframe(all_tests_df, use_container_width=True, hide_index=True)
                student_topic_rows = [row for row in topic_progress["by_student"] if row["Student"] == selected_student]
                if student_topic_rows:
                    st.markdown("**Topic progress for selected student**")
                    st.dataframe(pd.DataFrame(student_topic_rows), use_container_width=True, hide_index=True)

    if aggregate["variant_comparison"]:
        st.markdown("**Variant comparison A/B/C/D**")
        st.dataframe(pd.DataFrame(aggregate["variant_comparison"]), use_container_width=True, hide_index=True)

    if aggregate["timeline"]:
        st.markdown("**Performance over time**")
        timeline_df = pd.DataFrame(aggregate["timeline"]).set_index("Date")
        st.line_chart(timeline_df[["Average %"]])

    st.markdown("**Hardest Questions**")
    if aggregate["question_insights"]:
        st.dataframe(pd.DataFrame(aggregate["question_insights"][:10]), use_container_width=True, hide_index=True)
    else:
        st.info("Question-level insights will appear after submissions.")

    st.markdown("**Recent attempts**")
    attempt_df = pd.DataFrame(
        [
            {
                "ID": item["id"],
                "Student": item["student_name"],
                "Variant": item["variant_name"],
                "Test": item["test_title"],
                "Share Token": item.get("share_token", ""),
                "Score %": item["percentage"],
                "Risk": classify_risk(float(item["percentage"])),
                "Date": item["created_at"],
            }
            for item in attempts
        ]
    )
    st.dataframe(attempt_df, use_container_width=True, hide_index=True)


@st.fragment(run_every="5s")
def render_live_analytics_panel() -> None:
    """Auto-refresh analytics so new student submissions appear quickly."""
    st.caption("Live updates: this analytics view refreshes automatically every 5 seconds.")
    render_analytics_dashboard()


def render_student_journal(attempts: list[dict[str, Any]]) -> None:
    """Render a compact journal for the current test."""
    if not attempts:
        st.info("No journal entries yet.")
        return
    journal_rows = []
    by_student: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for item in attempts:
        by_student[item["student_name"]].append(item)
    for student_name, student_attempts in by_student.items():
        percentages = [float(item["percentage"]) for item in student_attempts]
        journal_rows.append(
            {
                "Student": student_name,
                "Attempts": len(student_attempts),
                "Latest %": round(percentages[0], 2),
                "Average %": round(sum(percentages) / len(percentages), 2),
                "Best %": round(max(percentages), 2),
                "Risk": classify_risk(sum(percentages) / len(percentages)),
            }
        )
    journal_rows.sort(key=lambda row: row["Average %"])
    st.dataframe(pd.DataFrame(journal_rows), use_container_width=True, hide_index=True)


def render_attempt_admin_tools(selected_attempt: dict[str, Any]) -> None:
    """Allow teachers to review, edit, or delete one saved attempt."""
    with st.expander("Review or edit this attempt", expanded=False):
        with st.form(f"attempt_review_form_{selected_attempt['id']}"):
            edited_name = st.text_input("Student name", value=selected_attempt.get("student_name", ""))
            edited_percentage = st.number_input(
                "Score %",
                min_value=0.0,
                max_value=100.0,
                value=float(selected_attempt.get("percentage", 0.0)),
                step=0.5,
            )
            edited_status = st.selectbox(
                "Review status",
                options=["submitted", "reviewed", "flagged", "excused"],
                index=["submitted", "reviewed", "flagged", "excused"].index(
                    str(selected_attempt.get("review_status", "submitted")) if str(selected_attempt.get("review_status", "submitted")) in {"submitted", "reviewed", "flagged", "excused"} else "submitted"
                ),
            )
            teacher_note = st.text_area(
                "Teacher note",
                value=str(selected_attempt.get("teacher_note", "")),
                height=100,
            )
            save_review = st.form_submit_button("Save changes", use_container_width=True)
        if save_review:
            ok = update_attempt_result(
                attempt_id=int(selected_attempt["id"]),
                student_name=edited_name,
                percentage=edited_percentage,
                review_status=edited_status,
                teacher_note=teacher_note,
            )
            if ok:
                log_event("update_attempt", "attempt", str(selected_attempt["id"]), {"student_name": edited_name, "percentage": edited_percentage, "review_status": edited_status})
                st.success("Attempt updated.")
                st.rerun()
            st.error("Attempt could not be updated.")

        if st.button("Delete this attempt", key=f"delete_attempt_{selected_attempt['id']}", use_container_width=True):
            if delete_attempt_result(int(selected_attempt["id"])):
                log_event("delete_attempt", "attempt", str(selected_attempt["id"]), {"student_name": selected_attempt.get("student_name", "")})
                st.success("Attempt deleted.")
                st.rerun()
            st.error("Attempt could not be deleted.")


def render_student_answers_view() -> None:
    """Render a detailed on-site view of saved student answers."""
    st.subheader("Student Answers")
    attempts = list_attempt_results(limit=200, owner_email=get_owner_email(), test_uid=get_current_test_uid())
    st.caption("This table is scoped to the current test only.")
    if not attempts:
        st.info("No student answers have been submitted yet for this test.")
        return

    student_options = ["All"] + sorted({item["student_name"] for item in attempts})
    test_options = ["All"] + sorted({item["test_title"] for item in attempts})
    variant_options = ["All"] + sorted({item["variant_name"] for item in attempts})

    filter_col1, filter_col2, filter_col3 = st.columns(3, gap="large")
    with filter_col1:
        selected_student = st.selectbox("Filter by student", options=student_options, key="answers_student_filter")
    with filter_col2:
        selected_test = st.selectbox("Filter by test", options=test_options, key="answers_test_filter")
    with filter_col3:
        selected_variant = st.selectbox("Filter by variant", options=variant_options, key="answers_variant_filter")

    filtered_attempts = attempts
    if selected_student != "All":
        filtered_attempts = [item for item in filtered_attempts if item["student_name"] == selected_student]
    if selected_test != "All":
        filtered_attempts = [item for item in filtered_attempts if item["test_title"] == selected_test]
    if selected_variant != "All":
        filtered_attempts = [item for item in filtered_attempts if item["variant_name"] == selected_variant]

    if not filtered_attempts:
        st.warning("No attempts match the selected filters.")
        return

    with st.expander("Student journal", expanded=False):
        render_student_journal(filtered_attempts)

    summary_df = pd.DataFrame(
        [
            {
                "ID": item["id"],
                "Student": item["student_name"],
                "Test": item["test_title"],
                "Variant": item["variant_name"],
                "Share Token": item.get("share_token", ""),
                "Score %": item["percentage"],
                "Status": item.get("review_status", "submitted"),
                "Risk": classify_risk(float(item["percentage"])),
                "Submitted At": item["created_at"],
            }
            for item in filtered_attempts
        ]
    )
    st.dataframe(summary_df, use_container_width=True, hide_index=True)

    attempt_ids = [item["id"] for item in filtered_attempts]
    selected_attempt_id = st.selectbox(
        "Open attempt details",
        options=attempt_ids,
        format_func=lambda attempt_id: next(
            (
                f"#{item['id']} | {item['student_name']} | {item['test_title']} | {item['percentage']}%"
                for item in filtered_attempts
                if item["id"] == attempt_id
            ),
            str(attempt_id),
        ),
        key="answers_attempt_selector",
    )

    selected_attempt = next(item for item in filtered_attempts if item["id"] == selected_attempt_id)
    details = selected_attempt["details"]

    metric_col1, metric_col2, metric_col3, metric_col4 = st.columns(4)
    metric_col1.metric("Student", selected_attempt["student_name"])
    metric_col2.metric("Variant", selected_attempt["variant_name"])
    metric_col3.metric("Score", f"{selected_attempt['percentage']}%")
    metric_col4.metric("Risk", classify_risk(float(selected_attempt["percentage"])))
    if selected_attempt.get("teacher_note"):
        st.info(f"Teacher note: {selected_attempt['teacher_note']}")
    render_attempt_admin_tools(selected_attempt)

    st.markdown("**Submitted answers**")
    for item in details.get("per_question", []):
        with st.container(border=True):
            st.markdown(f"**Question {item['index']}**")
            st.write(item["question"])
            st.caption(f"Skill: {item.get('skill_tag', '')}")
            st.write(f"Student answer: {item.get('student_answer', '')}")
            if item.get("correct_answer"):
                st.write(f"Correct answer: {item['correct_answer']}")
            if item.get("explanation"):
                st.write(f"Explanation: {item['explanation']}")
            st.write(f"Score for this question: {item.get('score', 0)}")


@st.fragment(run_every="5s")
def render_live_answers_panel() -> None:
    """Auto-refresh student answers so teacher view updates shortly after submission."""
    st.caption("Live updates: new student answers appear automatically every 5 seconds.")
    render_student_answers_view()


def load_history_item(record_id: int) -> None:
    """Load a saved test snapshot into the current editor."""
    payload = load_test_record(record_id)
    if payload is None:
        st.sidebar.error("The selected history record could not be loaded.")
        return
    load_test_payload(payload)
    st.sidebar.success(f"Loaded history record #{record_id}.")
    st.rerun()


def load_test_by_uid(test_uid: str) -> None:
    """Load the latest test snapshot by test UID."""
    payload = load_latest_test_record(test_uid, get_owner_email())
    if payload is None:
        st.error("The selected test could not be loaded.")
        return
    load_test_payload(payload)
    st.success("Test loaded into the workspace.")
    st.rerun()


def load_bank_item_into_test(record_id: int) -> None:
    """Append a question bank item to the current test."""
    payload = load_question_bank_item(record_id)
    if payload is None:
        st.sidebar.error("The selected question bank item could not be loaded.")
        return
    if st.session_state.generated_test is None:
        st.sidebar.error("Generate or load a test first.")
        return

    st.session_state.generated_test["questions"].append(payload)
    st.session_state.editor_version += 1
    st.session_state.quality_report = analyze_test_quality(st.session_state.generated_test)
    st.sidebar.success(f"Added question bank item #{record_id} to the current test.")
    st.rerun()


def render_history_sidebar() -> None:
    """Render recent saved tests from SQLite history."""
    with st.sidebar.expander("History", expanded=False):
        notice = st.session_state.get("history_notice", "")
        if notice:
            st.caption(notice)

        query = st.text_input("Search history", placeholder="Search by title or topic")
        history_items = list_test_library(owner_email=get_owner_email(), include_archived=False)
        if query.strip():
            query_lower = query.lower().strip()
            history_items = [
                item
                for item in history_items
                if query_lower in item["title"].lower()
                or query_lower in item["topic"].lower()
                or query_lower in str(item.get("source_name", "")).lower()
            ]

        if not history_items:
            st.info("No saved tests yet.")
        else:
            for item in history_items:
                with st.container(border=True):
                    st.markdown(f"**{item['title']}**")
                    st.caption(
                        f"{item['updated_at']} | {item['test_type']} | "
                        f"{item['difficulty']} | {item['language']}"
                    )
                    if item.get("grade_level") or item.get("assessment_purpose"):
                        st.caption(
                            f"Grade: {item.get('grade_level', '')} | Purpose: {item.get('assessment_purpose', '')}"
                        )
                    if item.get("source_name"):
                        st.caption(f"Source: {item['source_name']}")
                    if st.button("Load", key=f"history_load_{item['test_uid']}", use_container_width=True):
                        load_test_by_uid(item["test_uid"])


def render_test_library_view() -> None:
    """Render a cleaner library view for saved teacher tests."""
    open_section("Test Library")
    filter_col1, filter_col2, filter_col3, filter_col4, filter_col5, filter_col6 = st.columns(6, gap="large")
    with filter_col1:
        search = st.text_input("Search", placeholder="Title or source", key="library_search")
    with filter_col2:
        language_filter = st.selectbox(
            "Language",
            options=["All"] + list(LANGUAGE_OPTIONS),
            format_func=lambda key: "All" if key == "All" else LANGUAGE_OPTIONS[key],
            key="library_language",
        )
    with filter_col3:
        grade_filter = st.selectbox(
            "Grade",
            options=["All"] + GRADE_LEVEL_OPTIONS,
            key="library_grade",
        )
    with filter_col4:
        topic_filter = st.text_input("Topic filter", placeholder="Topic", key="library_topic")
    with filter_col5:
        subject_tag_filter = st.text_input("Subject tag", placeholder="Tag", key="library_subject_tag")
    with filter_col6:
        date_filter = st.text_input("Updated date", placeholder="YYYY-MM-DD", key="library_date")

    toggle_col1, toggle_col2, toggle_col3 = st.columns(3, gap="large")
    with toggle_col1:
        include_archived = st.checkbox("Show archived tests", value=False, key="library_archived")
    with toggle_col2:
        favorites_only = st.checkbox("Favorites only", value=False, key="library_favorites")
    with toggle_col3:
        sort_by = st.selectbox(
            "Sort by",
            options=["updated_desc", "updated_asc", "grade", "language", "title"],
            format_func=lambda value: {
                "updated_desc": "Newest first",
                "updated_asc": "Oldest first",
                "grade": "Grade",
                "language": "Language",
                "title": "Title",
            }[value],
            key="library_sort",
        )

    items = list_test_library(
        owner_email=get_owner_email(),
        search=search,
        language="" if language_filter == "All" else language_filter,
        grade_level="" if grade_filter == "All" else grade_filter,
        topic=topic_filter,
        subject_tag=subject_tag_filter,
        include_archived=include_archived,
        favorites_only=favorites_only,
        sort_by=sort_by,
    )
    if date_filter.strip():
        items = [item for item in items if str(item.get("updated_at", "")).startswith(date_filter.strip())]
    if not items:
        st.info("No saved tests match the selected filters.")
        close_section()
        return

    selected_test_uids = st.multiselect(
        "Bulk-select tests",
        options=[item["test_uid"] for item in items],
        format_func=lambda uid: next((f"{item['title']} | {item['grade_level']} | {item['updated_at']}" for item in items if item["test_uid"] == uid), uid),
        key="library_bulk_tests",
    )
    bulk_col1, bulk_col2 = st.columns(2, gap="large")
    with bulk_col1:
        if selected_test_uids and st.button("Archive selected tests", use_container_width=True):
            for uid in selected_test_uids:
                set_test_archived(uid, get_owner_email(), True)
            log_event("bulk_archive_tests", "test", ",".join(selected_test_uids), {"count": len(selected_test_uids)})
            st.rerun()
    with bulk_col2:
        if selected_test_uids:
            export_rows = [item for item in items if item["test_uid"] in selected_test_uids]
            st.download_button(
                "Export selected tests CSV",
                pd.DataFrame(export_rows).to_csv(index=False).encode("utf-8"),
                "selected_tests.csv",
                "text/csv",
                use_container_width=True,
            )

    for item in items:
        with st.container(border=True):
            title_col, action_col = st.columns([4, 1], gap="large")
            with title_col:
                st.markdown(f"**{item['title']}**")
                st.caption(
                    f"{item['topic']} | {item['language']} | {item['difficulty']} | "
                    f"{item['grade_level']} | Updated {item['updated_at']}"
                )
                if item.get("subject_tags"):
                    st.caption(f"Tags: {item['subject_tags']}")
                if item.get("source_name"):
                    st.caption(f"Source: {item['source_name']}")
                if int(item.get("archived", 0)):
                    st.caption("Archived")
                if int(item.get("is_favorite", 0)):
                    st.caption("Favorite")
            with action_col:
                if st.button("Load", key=f"library_load_{item['test_uid']}", use_container_width=True):
                    load_test_by_uid(item["test_uid"])
            more_col1, more_col2, more_col3, more_col4 = st.columns(4, gap="large")
            with more_col1:
                preview_key = f"library_preview_{item['test_uid']}"
                with st.popover("Preview", use_container_width=True):
                    payload = load_latest_test_record(item["test_uid"], get_owner_email())
                    if payload is None:
                        st.info("Preview unavailable.")
                    else:
                        preview_test = payload["test_data"]
                        st.markdown(f"**{preview_test.get('title', 'Generated Test')}**")
                        st.caption(
                            f"{len(preview_test.get('questions', []))} questions | "
                            f"{payload['metadata'].get('language', '')} | {payload['metadata'].get('grade_level', '')}"
                        )
                        for preview_index, question in enumerate(preview_test.get("questions", [])[:2], start=1):
                            st.write(f"{preview_index}. {question.get('question', '')}")
                        if len(preview_test.get("questions", [])) > 2:
                            st.caption("Open the test to view the full content.")
            with more_col2:
                if st.button("Duplicate", key=f"library_dup_{item['test_uid']}", use_container_width=True):
                    payload = load_latest_test_record(item["test_uid"], get_owner_email())
                    if payload is not None:
                        new_uid = uuid4().hex
                        payload["metadata"]["test_uid"] = new_uid
                        payload["test_data"]["test_uid"] = new_uid
                        payload["metadata"]["is_favorite"] = False
                        for variant_name, variant_data in payload.get("variants", {}).items():
                            variant_data["test_uid"] = new_uid
                            variant_data["variant_name"] = variant_name
                        load_test_payload(payload)
                        record_id = save_current_test_snapshot(payload["test_data"], payload["metadata"])
                        st.session_state.history_notice = f"Duplicated as record #{record_id}."
                        st.rerun()
            with more_col3:
                favorite_label = "Unfavorite" if int(item.get("is_favorite", 0)) else "Favorite"
                if st.button(favorite_label, key=f"library_favorite_{item['test_uid']}", use_container_width=True):
                    set_test_favorite(item["test_uid"], get_owner_email(), not bool(int(item.get("is_favorite", 0))))
                    if get_current_test_uid() == item["test_uid"]:
                        st.session_state.test_metadata["is_favorite"] = not bool(int(item.get("is_favorite", 0)))
                    st.rerun()
            with more_col4:
                archive_label = "Unarchive" if int(item.get("archived", 0)) else "Archive"
                if st.button(archive_label, key=f"library_archive_{item['test_uid']}", use_container_width=True):
                    set_test_archived(item["test_uid"], get_owner_email(), not bool(int(item.get("archived", 0))))
                    st.rerun()
    close_section()


def render_question_bank_sidebar() -> None:
    """Render question bank controls."""
    with st.sidebar.expander("Question Bank", expanded=False):
        notice = st.session_state.get("question_bank_notice", "")
        if notice:
            st.caption(notice)

        search = st.text_input("Search bank", placeholder="Search questions", key="bank_search")
        items = list_question_bank(limit=30, owner_email=get_owner_email())
        if search.strip():
            search_lower = search.lower().strip()
            items = [
                item
                for item in items
                if search_lower in item["question_text"].lower()
                or search_lower in item["topic"].lower()
                or search_lower in item["skill_tag"].lower()
            ]

        if not items:
            st.info("No saved questions yet.")
            return

        for item in items:
            with st.container(border=True):
                st.markdown(f"**#{item['id']} {item['question_type']}**")
                st.caption(item["question_text"][:110] + ("..." if len(item["question_text"]) > 110 else ""))
                st.caption(f"Topic: {item.get('topic', '')} | Skill: {item.get('skill_tag', '')}")
                if st.button("Add To Test", key=f"bank_add_{item['id']}", use_container_width=True):
                    load_bank_item_into_test(int(item["id"]))


def build_demo_responses(variant_data: dict[str, Any], mode: str = "medium") -> dict[str, Any]:
    """Build deterministic demo responses for defense-ready seeded attempts."""
    responses: dict[str, Any] = {}
    for index, question in enumerate(variant_data.get("questions", [])):
        if question["type"] in {"multiple_choice", "true_false"}:
            options = question.get("options", [])
            if mode == "strong":
                responses[f"question_{index}"] = question.get("correct_answer", "")
            elif mode == "weak":
                responses[f"question_{index}"] = options[-1] if options else ""
            else:
                responses[f"question_{index}"] = question.get("correct_answer", "") if index % 2 == 0 else (options[-1] if options else "")
        elif question["type"] == "short_answer":
            responses[f"question_{index}"] = question.get("correct_answer", "") if mode != "weak" else "I am not sure"
        elif question["type"] == "matching":
            pairs = question.get("pairs", [])
            if mode == "strong":
                responses[f"question_{index}"] = {pair["left"]: pair["right"] for pair in pairs}
            else:
                rotated = [pair["right"] for pair in pairs[1:]] + [pairs[0]["right"]] if pairs else []
                responses[f"question_{index}"] = {
                    pair["left"]: (pair["right"] if mode == "medium" and pair_index % 2 == 0 else rotated[pair_index] if pair_index < len(rotated) else "")
                    for pair_index, pair in enumerate(pairs)
                }
    return responses


def prepare_demo_accounts_and_attempts() -> str:
    """Create defense-friendly local demo profiles and sample attempts."""
    messages = []
    for email, password, display_name, role in (
        ("teacher.demo@local", "TeacherDemo2026!", "Demo Teacher", "teacher"),
        ("student.demo@local", "StudentDemo2026!", "Demo Student", "student"),
    ):
        ok, message = create_local_user(email, password, display_name, role)
        if ok:
            messages.append(f"Created {role} demo account: {email}")
        elif "already exists" in message.lower():
            messages.append(f"Demo account already exists: {email}")

    current_test = st.session_state.get("generated_test")
    if current_test and get_current_test_uid():
        variants = get_effective_variants(current_test)
        seed_plan = [
            ("Demo Student", "Variant A", "strong"),
            ("Demo Student", "Variant B", "medium"),
            ("Demo Student", "Variant D", "weak"),
        ]
        for offset, (student_name, variant_name, mode) in enumerate(seed_plan, start=1):
            variant_data = variants[variant_name]
            responses = build_demo_responses(variant_data, mode=mode)
            result = grade_attempt(variant_data, responses)
            result["responses"] = responses
            submission_key = hashlib.sha256(
                f"demo::{get_current_test_uid()}::{variant_name}::{offset}".encode("utf-8")
            ).hexdigest()
            if attempt_submission_exists(submission_key):
                continue
            save_attempt_result(
                student_name=student_name,
                test_uid=get_current_test_uid(),
                variant_name=variant_name,
                test_title=variant_data.get("title", ""),
                owner_email=get_owner_email(),
                share_token="demo-seed",
                submission_key=submission_key,
                percentage=result["percentage"],
                payload=result,
            )
        messages.append("Seeded demo attempts for the current test.")
    else:
        messages.append("Demo accounts are ready. Generate one test first if you also want seeded attempts.")

    return " ".join(messages)


def render_defense_materials_notice() -> None:
    """Show where the user can find defense materials."""
    with st.sidebar.expander("Defense Materials", expanded=False):
        st.markdown(
            """
Use these local files during your project defense:

- `README.md`
- `PROJECT_DEFENSE_KZ.md`
- `CRITERIA_MAP_KZ.md`
- `supabase_schema.sql`
- `DEMO_SCRIPT.md`

The suggested demo order is:
1. Generate a test from a topic
2. Generate another from a file
3. Edit one question and save it to the bank
4. Show variants A/B/C/D
5. Run student mode
6. Open analytics dashboard
"""
        )
        if st.button("Prepare demo accounts and attempts", use_container_width=True):
            st.success(prepare_demo_accounts_and_attempts())
        logs = list_api_error_logs(limit=5)
        if logs:
            with st.expander("Recent API errors", expanded=False):
                for item in logs:
                    st.caption(f"{item['created_at']} | {item['provider']}")
                    st.write(item["error_message"])



def render_empty_workspace_state() -> None:
    """Render a polished placeholder before the first generation."""
    st.markdown(
        """
        <div class="workspace-hint">
            <strong>No active test yet.</strong><br/>
            Start in Create. After generation, the interface will open a staged workflow:
            Review → Share → Analyze.
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_output() -> None:
    """Render the editable generated test and advanced project features."""
    generated_test = st.session_state.generated_test
    if not generated_test:
        render_empty_workspace_state()
        return

    metadata = st.session_state.test_metadata
    render_test_status_banner()
    action_col1, action_col2 = st.columns([5, 1], gap="large")
    with action_col1:
        st.caption("Flow: Create → Review → Share → Analyze")
    with action_col2:
        render_quick_actions()

    main_tabs = st.tabs(["Review", "Share", "Analyze"])

    with main_tabs[0]:
        st.session_state.active_flow_step = "Review"
        review_tabs = st.tabs(["Snapshot", "Edit"])
        with review_tabs[0]:
            render_metadata_summary()
            render_source_preview(generated_test)
            if st.session_state.quality_report:
                with st.expander("Quality Summary", expanded=True):
                    render_quality_report(st.session_state.quality_report)
            else:
                quality_report = analyze_test_quality(
                    generated_test,
                    expected_count=metadata.get("question_count"),
                )
                st.session_state.quality_report = quality_report
                with st.expander("Quality Summary", expanded=True):
                    render_quality_report(quality_report)

        with review_tabs[1]:
            edited_test = render_test_editor(generated_test)
            st.session_state.generated_test = edited_test
            for key in (
                "topic",
                "language",
                "test_type",
                "grade_level",
                "learning_objective",
                "lesson_stage",
                "assessment_purpose",
                "source_summary",
                "key_concepts",
            ):
                if key in metadata or key in edited_test:
                    edited_test[key] = metadata.get(key, edited_test.get(key))

            quality_report = analyze_test_quality(
                edited_test,
                expected_count=metadata.get("question_count"),
            )
            st.session_state.quality_report = quality_report
            autosave_current_test(edited_test, st.session_state.test_metadata)
            editor_variant_name = metadata.get("editor_variant_name", "Variant D")
            st.caption(f"Current editor view: {editor_variant_name}")
            render_save_snapshot_button(edited_test)

    latest_test = st.session_state.generated_test
    latest_quality_report = analyze_test_quality(
        latest_test,
        expected_count=metadata.get("question_count"),
    )
    st.session_state.quality_report = latest_quality_report
    variants = get_effective_variants(latest_test)

    with main_tabs[1]:
        st.session_state.active_flow_step = "Share"
        share_tabs = st.tabs(["Variants & Export", "Practice"])
        with share_tabs[0]:
            render_variants_section(latest_test, disable=not latest_quality_report["is_export_ready"])
        with share_tabs[1]:
            render_student_mode(variants, disable=not latest_quality_report["is_export_ready"])

    with main_tabs[2]:
        st.session_state.active_flow_step = "Analyze"
        analyze_tabs = st.tabs(["Home", "Dashboard", "Gradebook", "Answers", "Roster", "Library", "Backup", "Business"])
        with analyze_tabs[0]:
            render_teacher_home_dashboard()
        with analyze_tabs[1]:
            render_live_analytics_panel()
        with analyze_tabs[2]:
            render_gradebook_view()
        with analyze_tabs[3]:
            render_live_answers_panel()
        with analyze_tabs[4]:
            render_groups_and_roster_view()
        with analyze_tabs[5]:
            render_test_library_view()
        with analyze_tabs[6]:
            render_backup_center()
        with analyze_tabs[7]:
            render_business_center()


def main() -> None:
    """Run the Streamlit app."""
    initialize_state()
    render_theme()
    share_token = get_share_token_from_query()
    if share_token:
        render_shared_student_page(share_token)
        return

    render_profile_sidebar()
    render_cloud_status_sidebar()
    render_history_sidebar()
    render_question_bank_sidebar()
    render_share_links_sidebar()

    render_header()
    render_onboarding_panel()
    render_project_explainers()
    render_defense_materials_notice()

    if st.session_state.generated_test is None:
        shell_tabs = st.tabs(["Create", "Library"])
        with shell_tabs[0]:
            form_data, generate_clicked = render_generator_form()
            if generate_clicked:
                handle_generation(form_data)
        with shell_tabs[1]:
            render_test_library_view()
            render_empty_workspace_state()
    else:
        shell_tabs = st.tabs(["Create", "Workspace"])
        with shell_tabs[0]:
            form_data, generate_clicked = render_generator_form()
            if generate_clicked:
                handle_generation(form_data)
        with shell_tabs[1]:
            render_output()


if __name__ == "__main__":
    main()
